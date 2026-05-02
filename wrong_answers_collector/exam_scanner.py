#!/usr/bin/env python3
"""
试卷截图错题整理脚本
使用OCR识别试卷截图中的错题，并生成Word文档。

依赖库：
- pix2text (Pix2Text) - 先进的OCR工具，支持文字、数学公式、表格识别
- python-docx
- Pillow (PIL)

安装命令：
pip install pix2text python-docx Pillow

更新说明：
- 2024年: 从 PaddleOCR 升级到 Pix2Text
- 优势: 更好的数学公式识别、表格识别、混合内容处理
- 优化: 修复乱码问题，添加文本清洗，Pix2Text 单例模式
"""

import os
import sys
# 禁用网络检查，使用本地模型缓存
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'
# 强制 UTF-8 编码输出，避免控制台乱码
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import argparse
import re
import unicodedata
from pathlib import Path
import cv2
import numpy as np
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pix2text import Pix2Text

# ---- 集成 Doc-Image-Tool ----
_DOC_TOOL_DIR = Path(__file__).parent.parent / 'third-party-libs' / 'Doc-Image-Tool' / 'function_method'
if str(_DOC_TOOL_DIR) not in sys.path:
    sys.path.insert(0, str(_DOC_TOOL_DIR))

# 导入 Doc-Image-Tool 的笔记去噪美化模块
from HandwritingDenoisingBeautifying import (
    docscan_main,
    sample_pixels,
    get_palette,
    apply_palette,
)
# 导入 Doc-Image-Tool 的漂白模块
from DocBleach import sauvola_threshold

# ============================================================
# 正则表达式定义
# ============================================================

# 题号前缀匹配 - 更全面地识别各种题号格式
# 支持: 1. 1、 1) (1) ① 1. 【1】 【第1题】 第1题 等
QUESTION_PREFIX_RE = re.compile(
    r'^\s*(?:'
    r'(?:第\s*)?(\d+)\s*[\.、\)\]]\s*'  # 1. 1、 1) 1] 等
    r'|(?:第\s*)?(\d+)\s*题\s*'          # 第1题 第 1 题
    r'|\(?(\d+)\)?\s*'                   # (1) 1
    r'|【(\d+)】\s*'                     # 【1】
    r'|①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩'         # 圆圈数字
    r'|\b([a-zA-Z])\s*[\.、\)]\s*'       # A. B、 C)
    r'|\b([a-zA-Z])\s*\)\s*'             # A) B)
    r')', re.IGNORECASE)

# 错题标记匹配 - 识别题干前后的各种错误标记
WRONG_MARKER_RE = re.compile(
    r'错(?!误)|'       # 错 (但不是错误)
    r'×|[xX]|'         # 乘号形式的错
    r'错误|'           # 错误
    r'答错|'
    r'不对|'
    r'不正确|'
    r'✗|✘|'            # 其他叉号
    r'❌|'             # emoji叉号
    r'\bwrong\b|'      # 英文wrong
    r'\bfalse\b|'      # 英文false
    r'\berror\b|'      # 英文error
    r'三角|'           # 三角形标记
    r'红色|'           # 红色标记
    r'红叉|'           # 红叉
    r'叉号|'           # 叉号
    r'打叉|'           # 打叉
    r'划掉|'           # 划掉
    r'删除|'           # 删除
    r'取消', re.IGNORECASE  # 取消
)

# ============================================================
# 图片级笔迹擦除（集成 Doc-Image-Tool）
# ============================================================

class _DenoiseOptions:
    """
    模拟 argparse.Namespace，供 docscan_main 使用。
    Doc-Image-Tool 的笔记去噪美化参数：
    - sample_fraction: 采样像素比例 (0~1)，默认 0.05
    - value_threshold: 背景亮度阈值 (0~1)，默认 0.25
    - sat_threshold: 背景饱和度阈值 (0~1)，默认 0.20
    - num_colors: K-means 聚类颜色数，默认 8
    - white_bg: 是否强制背景为白色
    - saturate: 是否对调色板做饱和度拉伸
    - quiet: 是否静默模式
    """
    def __init__(self, **kwargs):
        self.sample_fraction = kwargs.pop('sample_fraction', 0.05)
        self.value_threshold = kwargs.pop('value_threshold', 0.25)
        self.sat_threshold = kwargs.pop('sat_threshold', 0.20)
        self.num_colors = kwargs.pop('num_colors', 8)
        self.white_bg = kwargs.pop('white_bg', True)
        self.saturate = kwargs.pop('saturate', True)
        self.quiet = kwargs.pop('quiet', True)


def remove_handwriting_from_image(
    image_path: Path,
    output_path: Path,
    stroke_width_threshold: float = 2.5,
    solidity_threshold: float = 0.6,
) -> Path:
    """
    从试卷图片中擦除手写笔迹，保留印刷文字。
    
    处理管线：
    1. 先用 Doc-Image-Tool 的 Sauvola 阈值做文档漂白
    2. 再用 Doc-Image-Tool 的笔记去噪美化（K-Means 颜色量化）
    3. 最后用连通域分析 + 图像修复做精细擦除
    
    Args:
        image_path: 输入图片路径
        output_path: 输出图片路径
        stroke_width_threshold: 预留
        solidity_threshold: 预留
    
    Returns:
        输出图片路径
    """
    img = cv2.imread(str(image_path))
    if img is None:
        raise FileNotFoundError(f"无法读取图片: {image_path}")

    orig = img.copy()
    print(f"  原始尺寸: {orig.shape[1]}x{orig.shape[0]}")

    # ========== 阶段 1：文档漂白（Sauvola 自适应阈值）==========
    print("  阶段 1/3: 文档漂白（Sauvola）...")
    try:
        binary_bleach = sauvola_threshold(img, window_size=15, k=0.2, r=128)
        bleached = np.full_like(img, 255)
        text_mask = (binary_bleach == 0)
        bleached[text_mask] = orig[text_mask]
    except Exception as e:
        print(f"  漂白失败: {e}，跳过此阶段")
        bleached = orig

    # ========== 阶段 2：笔记去噪美化（K-Means 颜色量化）==========
    print("  阶段 2/3: 笔记去噪美化（K-Means 颜色量化）...")
    try:
        opts = _DenoiseOptions(
            num_colors=6, white_bg=True, saturate=True, quiet=True,
            sample_fraction=0.05, value_threshold=0.25, sat_threshold=0.20,
        )
        denoised = docscan_main(bleached, opts)
    except Exception as e:
        print(f"  去噪美化失败: {e}，跳过此阶段")
        denoised = bleached

    # ========== 阶段 3：精细擦除 ==========
    print("  阶段 3/3: 精细擦除...")
    try:
        # 在漂白后的原图上分析，产生掩膜
        erase_mask = _build_erase_mask(bleached)
        # 用白色填充（非 inpainting），避免污染印刷文字
        result = _white_fill_erase(denoised, erase_mask)
    except Exception as e:
        print(f"  精细擦除失败: {e}，使用阶段 2 结果")
        result = denoised

    cv2.imwrite(str(output_path), result)
    return output_path


def _build_erase_mask(bleached: np.ndarray) -> np.ndarray:
    """
    形态学开运算笔迹擦除掩膜。
    
    核心原理：形态学开运算自动区分粗细笔画。
    - 开运算(Opening) = 先腐蚀再膨胀 → 消除细笔画，保留粗笔画
    - binary - opened = 被消除的细笔画 = 手写笔迹候选
    
    然后保护印刷文字密集区：开运算后仍密集的区域 = 题干。
    """
    gray = cv2.cvtColor(bleached, cv2.COLOR_BGR2GRAY)
    gray = cv2.medianBlur(gray, 3)
    h, w = gray.shape

    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY_INV, blockSize=31, C=10
    )

    # ---- 形态学开运算：消除细笔画 ----
    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (2, 2))
    opened = cv2.morphologyEx(binary, cv2.MORPH_OPEN, k, iterations=1)

    # ---- 细笔画 = 原图 - 开运算 ----
    thin = cv2.subtract(binary, opened)

    # ---- 保护印刷文字密集区 ----
    # 开运算后仍保留的像素 → 粗笔画（印刷文字）
    # 对开运算结果做膨胀 → 找出印刷文字的"地盘"
    k_big = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (9, 9))
    print_zone = cv2.dilate(opened, k_big, iterations=1)
    # 印刷区内不擦除
    thin = cv2.bitwise_and(thin, cv2.bitwise_not(print_zone))

    # ---- 保护大块区域（示意图）----
    nl, lb, sts, _ = cv2.connectedComponentsWithStats(binary, connectivity=8)
    for i in range(1, nl):
        if sts[i, cv2.CC_STAT_AREA] > 600:
            thin[lb == i] = 0

    # ---- 形态学后处理 ----
    thin = cv2.morphologyEx(thin, cv2.MORPH_CLOSE, k, iterations=1)
    thin = cv2.morphologyEx(thin, cv2.MORPH_OPEN, k, iterations=1)
    thin = cv2.dilate(thin, k, iterations=1)

    thin_px = np.count_nonzero(thin)
    print(f"    开运算掩膜: {thin_px} px ({thin_px / (h * w) * 100:.1f}%)")
    return thin


def _white_fill_erase(img: np.ndarray, mask: np.ndarray) -> np.ndarray:
    """
    用白色填充擦除区域（替代 inpainting）。
    
    比 inpainting 更安全：不会把周围像素"涂抹"到文字区域造成污染。
    然后对擦除边缘做轻微模糊，过渡更自然。
    """
    if mask.max() == 0:
        return img
    
    result = img.copy()
    
    # 白色填充
    white = np.full_like(img, 255)
    result[mask > 0] = white[mask > 0]
    
    # 对擦除区域边缘做轻微高斯模糊，让过渡更自然
    edge_kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    mask_edges = cv2.dilate(mask, edge_kernel, iterations=1)
    mask_edges = cv2.subtract(mask_edges, mask)
    
    if mask_edges.max() > 0:
        blurred = cv2.GaussianBlur(result, (5, 5), 0)
        result[mask_edges > 0] = blurred[mask_edges > 0]
    
    return result


# ============================================================
# 文本级手写笔迹擦除（基于正则，OCR 后处理用）
# ============================================================

# 常见的手写答案提示词 - 这些后面的内容通常是手写答案
ANSWER_PREFIX_RE = re.compile(
    r'(?:答案|解答|解析|解|答|选|填|写)[:：]\s*|'  # 答案: 解答: 解: 答: 等
    r'(?:[Tt]herefore|[Bb]ecause|[Aa]nswer)[:：]\s*'  # 英文答案提示
)

# 手写填空标记：题目中的空白处，学生手写了答案
# 特征：括号/方框内的短内容，通常只有数字/字母/简单符号
HANDWRITING_FILL_RE = re.compile(
    r'(?:'
    # ( ) 或 （ ）中的手写填空（内容短，通常 ≤5 个字符，不含中文）
    r'[(（]\s*([^\u4e00-\u9fff()（）]{1,5})\s*[)）]'
    r'|'
    # 【 】中的手写填空
    r'【\s*([^\u4e00-\u9fff【】]{1,5})\s*】'
    r'|'
    # [ ] 中的手写填空
    r'\[\s*([^\u4e00-\u9fff\[\]]{1,5})\s*\]'
    r')'
)

# 明显是手写选择题答案的模式：单独成行的字母或字母组合
# 如 OCR 识别到孤立的一行 "A" "B" "C" "D" 或 "AB" "CD"
STRAY_CHOICE_RE = re.compile(
    r'^\s*[A-Da-d]{1,4}\s*$'
)

# 显著的手写标记：学生用星号、下划线等标记的内容
HANDWRITING_MARKER_RE = re.compile(
    r'\*[^*]{1,20}\*|'    # *手写标记内容*
    r'_{2,}|'             # 连续下划线（填空线）
    r'--{2,}|'            # 连续横线
    r'~~|'                # 删除线标记
    r'√|'                 # 对勾
    r'[/\\]'              # 随手画的斜线
)


def erase_handwriting(text: str) -> str:
    """
    擦除 OCR 识别到的手写答题笔迹。
    
    试卷上学生手写的答案与印刷题目混合在一起，
    OCR 无法区分。此函数通过启发式规则移除常见的手写内容：
    
    1. 填空处的手写答案（括号/方框内短内容）
    2. 答案提示词后的手写内容（如"答：xxx"）
    3. 孤立的选择题答案字母（A/B/C/D）
    4. 手写标记符号（星号、下划线、对勾等）
    
    注意：会用占位符替代被擦除的内容，保留题目结构。
    """
    if not text:
        return text
    
    # ---- 1. 移除填空括号内的手写答案 ----
    # 保留括号本身，只清除内容，避免破坏题目结构
    def _clean_fill(match):
        full = match.group(0)
        # 检查内容是否像手写答案（不含中文、不含复杂标点）
        inner = (match.group(1) or match.group(2) or match.group(3) or '').strip()
        if not inner:
            return full  # 空括号，保留
        # 如果内容包含中文，大概率是印刷内容，保留
        if re.search(r'[\u4e00-\u9fff]', inner):
            return full
        # 手写答案 — 保留括号结构，用空格替代内容
        if full.startswith('(') or full.startswith('（'):
            return '( )' if full.startswith('(') else '（ ）'
        elif full.startswith('【'):
            return '【 】'
        elif full.startswith('['):
            return '[ ]'
        return full
    
    text = HANDWRITING_FILL_RE.sub(_clean_fill, text)
    
    # ---- 2. 移除答案提示词后的手写内容 ----
    # 如 "答：3.14" → 移除 "3.14"，保留 "答：___"
    def _clean_answer_prefix(match):
        return match.group(0)  # 保留提示词本身
    
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        # 如果行以答案提示词开头，尝试截断手写部分
        m = ANSWER_PREFIX_RE.match(line)
        if m:
            prefix = m.group(0)
            rest = line[m.end():].strip()
            if rest:
                # 检查剩余部分是否像手写（短、无句号、无中文长句）
                if (len(rest) <= 15 and 
                    not re.search(r'[。！？]', rest) and
                    not re.search(r'[\u4e00-\u9fff]{4,}', rest)):
                    # 疑似手写答案，保留提示词，用占位符替代答案
                    cleaned_lines.append(prefix + '[手写答案已擦除]')
                    continue
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # ---- 3. 移除孤立的选择题答案字母 ----
    lines = text.splitlines()
    cleaned_lines = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if STRAY_CHOICE_RE.match(stripped):
            # 检查上下文：前后行是否有题目内容
            prev_is_content = (i > 0 and len(lines[i-1].strip()) > 10)
            next_is_content = (i < len(lines)-1 and len(lines[i+1].strip()) > 10)
            if prev_is_content or next_is_content:
                # 上下文有题目内容，这行很可能是手写答案 → 移除
                continue
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # ---- 4. 擦除手写标记符号 ----
    text = HANDWRITING_MARKER_RE.sub(' ', text)
    
    # ---- 5. 清理擦除后产生的多余空白 ----
    text = re.sub(r' {3,}', '  ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text

def clean_ocr_text(text: str) -> str:
    """
    清洗 OCR 识别结果，移除乱码和噪声字符。
    
    常见 OCR 乱码来源：
    - 控制字符 (U+0000-U+001F, U+007F-U+009F)
    - 零宽字符、双向控制字符
    - 私有区字符 (PUA)
    - 无法解码的替换字符 (U+FFFD)
    - 孤立的组合标记
    - OCR 产生的无意义符号序列
    """
    if not text:
        return text
    
    # 1. 移除替换字符 (U+FFFD) - 这是乱码的典型标志
    text = text.replace('\ufffd', '')
    
    # 2. 移除零宽字符和双向控制字符
    zero_width_and_bidi = set(
        chr(c) for c in range(0x200B, 0x200F + 1)  # 零宽空格等
    ) | set(
        chr(c) for c in range(0x2028, 0x202E + 1)  # 行分隔符、双向控制
    ) | set(
        chr(c) for c in range(0x2060, 0x206F + 1)  # 单词连接符、不可见字符
    ) | {
        '\ufeff',   # BOM / 零宽不换行空格
        '\ufff0',   # 其他替换字符
    }
    text = ''.join(ch for ch in text if ch not in zero_width_and_bidi)
    
    # 3. 移除其他控制字符（保留常用换行和制表符）
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # 4. 移除私有区字符 (PUA: U+E000-U+F8FF, U+F0000-U+FFFFF, U+100000-U+10FFFF)
    text = re.sub(r'[\ue000-\uf8ff\U000f0000-\U000fffff\U00100000-\U0010ffff]', '', text)
    
    # 5. 规范化 Unicode（NFC 组合形式，对中文更友好）
    text = unicodedata.normalize('NFC', text)
    
    # 6. 擦除手写答题笔迹（填空答案、选择项、答案提示后的手写内容等）
    text = erase_handwriting(text)
    
    # 7. 清理 OCR 常见噪声：连续重复标点、无意义符号序列
    text = re.sub(r'[。，、；：？！…]{4,}', '...', text)  # 过多标点合并
    text = re.sub(r'(?<![a-zA-Z0-9])[a-zA-Z]{1,2}(?![a-zA-Z0-9])', '', text)  # 孤立的1-2个英文字母
    text = re.sub(r'[^\S\n]{3,}', '  ', text)  # 多个空格合并
    
    # 8. 清理 OCR 中常见的错误识别模式
    # 中文被识别为乱码拼音组合
    text = re.sub(r'[a-z]{15,}', '', text, flags=re.IGNORECASE)  # 超长无意义字母串
    
    # 9. 移除垃圾行（伪中文乱码、无意义符号行、OCR 幻想文本）
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append('')
            continue
        
        ratio = _line_valid_char_ratio(stripped)
        
        # 规则 A: 有效字符比例太低 → 垃圾
        if len(stripped) >= 5 and ratio < 0.35:
            continue
        
        # 规则 B: 短行（5-25 字符）全是符号/数字/乱码 → 垃圾
        if 5 <= len(stripped) <= 25 and ratio < 0.5:
            continue
        
        # 规则 C: 包含大量 CJK 字符但不通顺 → 伪中文乱码
        cjk_count = sum(1 for ch in stripped if '\u4e00' <= ch <= '\u9fff')
        if cjk_count >= 10 and not is_coherent_chinese(stripped):
            continue
        
        # 规则 D: 超长无意义行
        if len(stripped) > 80 and ratio < 0.6:
            continue
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def is_coherent_chinese(text: str) -> bool:
    """
    判断文本是否为通顺的中文（而非 OCR 产生的伪中文乱码）。
    
    原理：
    - 真实中文高频重用常见字（的、是、一、不、了...），
      字符重复率较高。
    - OCR 伪中文乱码通常由模型"幻想"出随机 CJK 字符，
      每个字都不同，重复率极低。
    
    Returns:
        True 如果文本看起来是通顺的中文，False 如果是伪中文乱码
    """
    if not text:
        return False
    
    # 提取所有 CJK 字符
    cjk_chars = [ch for ch in text if '\u4e00' <= ch <= '\u9fff']
    if len(cjk_chars) < 4:
        return True  # CJK 太少，不判断
    
    unique_cjk = len(set(cjk_chars))
    total_cjk = len(cjk_chars)
    
    # 唯一字符率：真实文本通常 < 0.7，乱码通常 > 0.85
    uniqueness = unique_cjk / total_cjk if total_cjk > 0 else 1.0
    
    if uniqueness > 0.85 and total_cjk >= 8:
        return False  # 几乎所有字都不同 → 伪中文
    
    if uniqueness > 0.75 and total_cjk >= 15:
        return False  # 长文本且唯一率高 → 可疑
    
    # 额外检查：真实中文通常包含高频字
    high_freq_chars = set('的是不了一有人在上个大这来为和国我以要到他会就出对生能而子说时下过得自开部家机可方后成所分前然没法如经其现当于从者并部度实定物权加量都两体制当计还资应关因重些特线')
    high_freq_count = sum(1 for ch in cjk_chars if ch in high_freq_chars)
    high_freq_ratio = high_freq_count / total_cjk if total_cjk > 0 else 0
    
    # 真实文本高频字占比通常 > 15%
    if total_cjk >= 10 and high_freq_ratio < 0.08:
        return False
    
    return True


def _line_valid_char_ratio(line: str) -> float:
    """计算一行中有效字符的比例"""
    if not line:
        return 0.0
    valid = sum(1 for ch in line if (
        '\u4e00' <= ch <= '\u9fff' or
        '\u3400' <= ch <= '\u4dbf' or
        '\uf900' <= ch <= '\ufaff' or
        ch.isalnum() or
        ch in ' .,;:!?()（）[]【】""''、。，；：？！…—–-+=*/<>≤≥≠≈±×÷√{}%#@&^|~'
    ))
    return valid / len(line) if line else 0.0


def is_garbled_text(text: str, threshold: float = 0.35) -> bool:
    """
    判断文本是否为乱码（综合检测）。
    
    结合字符比例分析和中文连贯性检测。
    """
    if not text or len(text) < 5:
        return False
    
    # 使用 _line_valid_char_ratio 做快速检查
    ratio = _line_valid_char_ratio(text)
    if ratio < threshold:
        return True
    
    # CJK 字符多但不通顺 → 伪中文乱码
    cjk_count = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
    if cjk_count >= 10 and not is_coherent_chinese(text):
        return True
    
    # 拉丁字母占比过高且无常见单词模式 → 乱码
    latin = sum(1 for ch in text if ch.isascii() and ch.isalpha())
    total = len(text)
    if latin > total * 0.8 and len(text) > 20:
        common_patterns = re.findall(r'\b[a-z]{2,8}\b', text.lower())
        if len(common_patterns) < len(text) * 0.1:
            return True
    
    return False


# 判断是否为题号行
def is_question_line(line):
    """判断行是否为题号/标记行"""
    # 完全匹配题号模式
    if QUESTION_PREFIX_RE.match(line):
        return True
    # 匹配纯题号（如只有 "1" 或 "(1)"）
    pure_question_re = re.compile(
        r'^\s*(\d+|[a-zA-Z]|①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩)\s*$'
    )
    if pure_question_re.match(line):
        return True
    return False

# 判断是否为错题标记行（只有标记或主要是标记）
def is_wrong_marker_heavy(line):
    """判断行是否主要是错题标记"""
    stripped = line.strip()
    if not stripped:
        return False
    # 匹配仅包含标记的模式
    marker_only_re = re.compile(
        r'^(?:错|×|x|错误|答错|不对|不正确|✗|✘|❌|wrong|error|三角|红色|红叉|叉号|打叉|划掉|删除|取消)\s*$', 
        re.IGNORECASE
    )
    return bool(marker_only_re.match(stripped))

# ============================================================
# Pix2Text 单例管理
# ============================================================

_p2t_instance = None

def get_p2t():
    """
    获取 Pix2Text 单例实例。
    只初始化一次，避免重复加载模型导致的性能问题和识别不稳定。
    """
    global _p2t_instance
    if _p2t_instance is None:
        print("正在加载 Pix2Text 模型（仅首次）...")
        try:
            _p2t_instance = Pix2Text(
                analyzer_config=dict(
                    text_detector_config=dict(
                        box_thresh=0.3,
                        unclip_ratio=2.0
                    )
                )
            )
            print("Pix2Text 模型加载完成。")
        except Exception as e:
            print(f"Pix2Text 模型加载失败: {e}")
            raise
    return _p2t_instance


def _extract_text_from_result(result) -> str:
    """
    从 Pix2Text 返回结果中递归提取纯文本。
    统一处理 dict / list / str 等多种返回格式，避免编码损失。
    """
    if isinstance(result, str):
        return result
    elif isinstance(result, dict):
        # 优先取已知字段
        for key in ('text', 'content', 'line_texts', 'sentences'):
            if key in result:
                val = result[key]
                if isinstance(val, str):
                    return val
                elif isinstance(val, list):
                    return '\n'.join(_extract_text_from_result(v) for v in val)
        # 遍历所有值，收集文本
        parts = []
        for val in result.values():
            extracted = _extract_text_from_result(val)
            if extracted:
                parts.append(extracted)
        return '\n'.join(parts)
    elif isinstance(result, (list, tuple)):
        return '\n'.join(
            _extract_text_from_result(item) for item in result
        )
    elif result is not None:
        return str(result)
    return ''


def extract_text_from_image(image_path: Path) -> str:
    """
    从图片中提取文本，并清洗 OCR 乱码。
    
    Args:
        image_path: 图片文件路径
        
    Returns:
        清洗后的文本字符串
    """
    try:
        p2t = get_p2t()
        result = p2t.recognize(str(image_path))
        raw_text = _extract_text_from_result(result)
        
        if not raw_text:
            print(f"  警告: 图片 {image_path.name} 未识别到文本")
            return ''
        
        # 清洗 OCR 噪声和乱码
        cleaned_text = clean_ocr_text(raw_text)
        
        # 检测是否为严重乱码（给出警告）
        if is_garbled_text(cleaned_text):
            print(f"  警告: 图片 {image_path.name} 识别结果疑似乱码，请检查图片质量")
        
        return cleaned_text
        
    except Exception as e:
        print(f"处理图片 {image_path} 时出错: {e}")
        return ""

def extract_question_number(line):
    """从行中提取题号"""
    match = QUESTION_PREFIX_RE.match(line)
    if match:
        # 查找第一个非空的捕获组（题号）
        for group in match.groups():
            if group:
                return group.strip()
    return None

def identify_wrong_answers(text):
    """
    识别错题：支持多种标记位置
    
    识别模式：
    1. 题号 + 错题标记（如：1. 错）
    2. 错题标记 + 题号 + 题干（如：错 1. xxx）
    3. 题号 + 包含标记的完整行（如：1. xxx 错）
    4. 标记在题目前的连续行
    """
    # 先清洗文本，去除 OCR 噪声
    text = clean_ocr_text(text)
    
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    wrong_questions = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 情况1: 题号后面跟着错题标记（标记在下一行）
        if is_question_line(line):
            question_num = extract_question_number(line)
            # 检查下一行是否是错题标记
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                if is_wrong_marker_heavy(next_line):
                    # 题号行 + 标记行
                    wrong_questions.append(f"{question_num} — {line}")
                    i += 2
                    continue
            # 检查当前行是否包含错题标记
            if WRONG_MARKER_RE.search(line):
                wrong_questions.append(f"{question_num} — {line}")
                i += 1
                continue
        
        # 情况2: 错题标记行，后面跟着题号
        if is_wrong_marker_heavy(line):
            # 检查是否紧跟着题号行
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                next_match = QUESTION_PREFIX_RE.match(next_line)
                if next_match:
                    question_num = extract_question_number(next_line)
                    wrong_questions.append(f"{question_num} — {next_line} (标记: {line})")
                    i += 2
                    continue
            # 单独一行错题标记，无法关联题号
            wrong_questions.append(f"未定位题号 — {line}")
            i += 1
            continue
        
        # 情况3: 当前行同时包含题号和错题标记
        if WRONG_MARKER_RE.search(line):
            question_num = extract_question_number(line)
            if question_num:
                wrong_questions.append(f"{question_num} — {line}")
            else:
                # 行中有标记但没有识别到题号
                wrong_questions.append(f"未定位题号 — {line}")
            i += 1
            continue
        
        i += 1
    
    # 去重（如果同一题被多次识别）
    seen = set()
    unique_wrong = []
    for q in wrong_questions:
        # 使用题号作为去重依据
        num_match = re.match(r'^(\d+|[a-zA-Z])', q)
        if num_match:
            num = num_match.group(1)
            if num not in seen:
                seen.add(num)
                unique_wrong.append(q)
        else:
            unique_wrong.append(q)
    
    return unique_wrong


def extract_all_questions(text):
    """提取所有题目块，按题号开始分组"""
    # 先清洗文本
    text = clean_ocr_text(text)
    
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    questions = []
    current_question = []

    for line in lines:
        if is_question_line(line):
            if current_question:
                questions.append(' '.join(current_question))
            current_question = [line]
        else:
            if current_question:
                current_question.append(line)

    if current_question:
        questions.append(' '.join(current_question))

    return questions


def create_word_document(items, output_path):
    """
    创建 Word 文档，设置中文字体避免打开时显示乱码。
    
    Args:
        items: 题目列表
        output_path: 输出文件路径
    """
    doc = Document()
    
    # ---- 设置默认字体为宋体，支持中文 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    # 设置中文字体（东亚字体）
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ---- 标题 ----
    heading = doc.add_heading('题目整理', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    if not items:
        p = doc.add_paragraph("未识别到题目。")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"共整理 {len(items)} 道题目", style='Normal')
        doc.add_paragraph('')  # 空行分隔
        
        for i, question in enumerate(items, 1):
            # 清理题目文本中可能残留的乱码
            clean_q = clean_ocr_text(question)
            
            heading = doc.add_heading(f'题目 {i}', level=1)
            for run in heading.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            para = doc.add_paragraph(clean_q)
            # 给每个题目段落也设置中文字体
            for run in para.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    try:
        doc.save(str(output_path))
        print(f"Word文档已保存到: {output_path}")
    except Exception as e:
        print(f"保存 Word 文档失败: {e}")
        raise

def main():
    parser = argparse.ArgumentParser(
        description='试卷截图处理工具：笔迹擦除 / OCR识别生成Word文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  # 只擦除笔迹，输出清洁图片
  python exam_scanner.py ./exam_screenshots --erase
  
  # OCR 识别所有题目，生成 Word 文档
  python exam_scanner.py ./exam_screenshots 错题集 --mode wrong
  python exam_scanner.py ./exam_screenshots 全部题目 --mode all
        '''
    )
    parser.add_argument('folder_path', help='包含试卷截图的文件夹路径')
    parser.add_argument('output_name', nargs='?', default=None,
                        help='生成的Word文档名称（不含扩展名），--erase 模式下不需要')
    parser.add_argument(
        '--erase', action='store_true',
        help='图片笔迹擦除模式：直接对图片做手写笔迹擦除，输出清洁图片（不生成Word文档）'
    )
    parser.add_argument(
        '--mode', choices=['all', 'wrong'], default='all',
        help='整理模式：all=整理所有题目，wrong=仅整理错题（默认: all）'
    )

    args = parser.parse_args()

    folder_path = Path(args.folder_path)
    if not folder_path.exists() or not folder_path.is_dir():
        print(f"错误: 文件夹路径不存在: {folder_path}")
        sys.exit(1)

    # 支持的图片格式
    image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp'}

    # 收集所有图片文件
    image_files = sorted([
        f for f in folder_path.iterdir()
        if f.suffix.lower() in image_extensions
    ])

    if not image_files:
        print(f"错误: 文件夹 {folder_path} 中未找到支持的图片文件")
        print(f"支持的格式: {', '.join(image_extensions)}")
        sys.exit(1)

    # ============================================================
    # 笔迹擦除模式
    # ============================================================
    if args.erase:
        output_dir = folder_path / "cleaned"
        output_dir.mkdir(exist_ok=True)

        print(f"笔迹擦除模式")
        print(f"找到 {len(image_files)} 张图片")
        print(f"输出目录: {output_dir}")
        print("-" * 50)

        success_count = 0
        for idx, file_path in enumerate(image_files, 1):
            output_path = output_dir / f"{file_path.stem}_cleaned{file_path.suffix}"
            print(f"[{idx}/{len(image_files)}] 擦除: {file_path.name}")

            try:
                remove_handwriting_from_image(file_path, output_path)
                print(f"  → 已保存: {output_path.name}")
                success_count += 1
            except Exception as e:
                print(f"  ✗ 失败: {e}")

        print("-" * 50)
        print(f"完成: {success_count}/{len(image_files)} 张图片已处理")
        return

    # ============================================================
    # OCR 识别模式
    # ============================================================
    if not args.output_name:
        print("错误: OCR 模式下需要提供 output_name 参数")
        print("示例: python exam_scanner.py ./screenshots 我的错题集 --mode wrong")
        sys.exit(1)

    output_name = args.output_name
    if output_name.lower().endswith('.docx'):
        output_name = output_name[:-5]
    output_path = folder_path / f"{output_name}.docx"

    print(f"找到 {len(image_files)} 张图片")
    print(f"整理模式: {'仅错题' if args.mode == 'wrong' else '所有题目'}")
    print("-" * 50)

    all_items = []
    success_count = 0

    for idx, file_path in enumerate(image_files, 1):
        print(f"[{idx}/{len(image_files)}] 处理: {file_path.name}")
        text = extract_text_from_image(file_path)

        if not text:
            print(f"  跳过: 未识别到文本")
            continue

        if args.mode == 'all':
            items = extract_all_questions(text)
        else:
            items = identify_wrong_answers(text)

        if items:
            all_items.extend(items)
            print(f"  识别到 {len(items)} 道题目")
        else:
            print(f"  未识别到题目")

        success_count += 1

    print("-" * 50)
    print(f"处理完成: {success_count}/{len(image_files)} 张图片成功识别")
    print(f"共整理 {len(all_items)} 道题目")

    create_word_document(all_items, output_path)

if __name__ == "__main__":
    main()
