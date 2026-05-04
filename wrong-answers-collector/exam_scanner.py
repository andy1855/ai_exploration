#!/usr/bin/env python3
"""
试卷截图错题整理脚本
使用OCR识别试卷截图中的错题，并生成Word文档。

依赖库：
- pix2text (Pix2Text) - 先进的OCR工具，支持文字、数学公式、表格识别
- python-docx
- Pillow (PIL)

安装命令：
pip install pix2text python-docx Pillow

更新说明：
- 2024年: 从 PaddleOCR 升级到 Pix2Text
- 优势: 更好的数学公式识别、表格识别、混合内容处理
- 优化: 修复乱码问题，添加文本清洗，Pix2Text 单例模式
"""

import os
import sys
# 禁用网络检查，使用本地模型缓存
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'
# 强制 UTF-8 编码输出，避免控制台乱码
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import argparse
import importlib.util
import shutil
import subprocess
import tempfile
import re
import unicodedata
from pathlib import Path
import cv2
import numpy as np
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pix2text import Pix2Text

# ---- 集成 Doc-Image-Tool ----
_DOC_TOOL_DIR = Path(__file__).parent.parent / 'third-party-libs' / 'Doc-Image-Tool' / 'function_method'
if str(_DOC_TOOL_DIR) not in sys.path:
    sys.path.insert(0, str(_DOC_TOOL_DIR))

# 导入 Doc-Image-Tool 的笔记去噪美化模块
from HandwritingDenoisingBeautifying import (
    docscan_main,
    sample_pixels,
    get_palette,
    apply_palette,
)
# 导入 Doc-Image-Tool 的漂白模块
from DocBleach import sauvola_threshold

# ============================================================
# 正则表达式定义
# ============================================================

# 题号前缀匹配 - 更全面地识别各种题号格式
# 支持: 1. 1、 1) (1) ① 1. 【1】 【第1题】 第1题 等
QUESTION_PREFIX_RE = re.compile(
    r'^\s*(?:'
    r'(?:第\s*)?(\d+)\s*[\.、\)\]]\s*'  # 1. 1、 1) 1] 等
    r'|(?:第\s*)?(\d+)\s*题\s*'          # 第1题 第 1 题
    r'|\(?(\d+)\)?\s*'                   # (1) 1
    r'|【(\d+)】\s*'                     # 【1】
    r'|①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩'         # 圆圈数字
    r'|\b([a-zA-Z])\s*[\.、\)]\s*'       # A. B、 C)
    r'|\b([a-zA-Z])\s*\)\s*'             # A) B)
    r')', re.IGNORECASE)

# 错题标记匹配 - 识别题干前后的各种错误标记
WRONG_MARKER_RE = re.compile(
    r'错(?!误)|'       # 错 (但不是错误)
    r'×|[xX]|'         # 乘号形式的错
    r'错误|'           # 错误
    r'答错|'
    r'不对|'
    r'不正确|'
    r'✗|✘|'            # 其他叉号
    r'❌|'             # emoji叉号
    r'\bwrong\b|'      # 英文wrong
    r'\bfalse\b|'      # 英文false
    r'\berror\b|'      # 英文error
    r'三角|'           # 三角形标记
    r'红色|'           # 红色标记
    r'红叉|'           # 红叉
    r'叉号|'           # 叉号
    r'打叉|'           # 打叉
    r'划掉|'           # 划掉
    r'删除|'           # 删除
    r'取消', re.IGNORECASE  # 取消
)

# ============================================================
# 图片级笔迹擦除（集成 Doc-Image-Tool）
# ============================================================

def _reference_image_stem(path: Path) -> str:
    """文件名主名 NFC 规范化，避免因全角括号等导致误判。"""
    return unicodedata.normalize('NFKC', path.stem).strip()


def _is_reference_only_image(path: Path) -> bool:
    """仅对照、不参与 `--erase` 批处理的参考图（与 raw 并排放时的效果图）。"""
    stem = _reference_image_stem(path).casefold()
    if '效果图' in stem or '效果圖' in stem:
        return True
    if stem == '效果图' or stem == '效果圖':
        return True
    if stem.startswith(('效果图-', '效果图_', '效果图 ')):
        return True
    if stem.startswith(('效果圖-', '效果圖_')):
        return True
    return False


class _DenoiseOptions:
    """
    模拟 argparse.Namespace，供 docscan_main 使用。
    Doc-Image-Tool 的笔记去噪美化参数：
    - sample_fraction: 采样像素比例 (0~1)，默认 0.05
    - value_threshold: 背景亮度阈值 (0~1)，默认 0.25
    - sat_threshold: 背景饱和度阈值 (0~1)，默认 0.20
    - num_colors: K-means 聚类颜色数，默认 8
    - white_bg: 是否强制背景为白色
    - saturate: 是否对调色板做饱和度拉伸
    - quiet: 是否静默模式
    """
    def __init__(self, **kwargs):
        self.sample_fraction = kwargs.pop('sample_fraction', 0.05)
        self.value_threshold = kwargs.pop('value_threshold', 0.25)
        self.sat_threshold = kwargs.pop('sat_threshold', 0.20)
        self.num_colors = kwargs.pop('num_colors', 8)
        self.white_bg = kwargs.pop('white_bg', True)
        self.saturate = kwargs.pop('saturate', True)
        self.quiet = kwargs.pop('quiet', True)


def remove_handwriting_from_image(
    image_path: Path,
    output_path: Path,
    erase_backend: str = 'doc-image-tool',
    model_repo: Path | None = None,
    model_checkpoint: Path | None = None,
    model_command: str | None = None,
    model_device: str = 'cpu',
    mask_threshold: float = 0.5,
) -> Path:
    """
    从 raw 拍屏图尽量接近扫描版效果图：去掉红批改、白化纸张、擦掉填空与草稿。

    1. HSV：仅高饱和鲜红色 → 白填（小膨胀）
    2. 亮度归一 + 拉伸；禁用锐化（避免白条变灰边）
    3. 左栏分列擦填空 + 稀疏行擦除 + 解答区上扩白填
    4. 左栏墨迹掩膜外交白、墨迹内压暗（消除灰雾与笔迹残影）

    Args:
        image_path: 输入图片路径
        output_path: 输出图片路径

    Returns:
        输出图片路径
    """
    img = cv2.imread(str(image_path))
    if img is None:
        raise FileNotFoundError(f"无法读取图片: {image_path}")

    h, w = img.shape[:2]
    print(f"  原始尺寸: {w}x{h}")

    if erase_backend != 'opencv':
        print(f"  模型擦除后端: {erase_backend}")
        try:
            result = remove_handwriting_with_model_backend(
                image_path=image_path,
                img=img,
                backend=erase_backend,
                model_repo=model_repo,
                model_checkpoint=model_checkpoint,
                model_command=model_command,
                model_device=model_device,
                mask_threshold=mask_threshold,
            )
        except Exception as e:
            print(f"  模型后端失败: {e}")
            if erase_backend != 'doc-image-tool':
                try:
                    print("  回退到 Doc-Image-Tool 管线。")
                    result = _run_doc_image_tool_pipeline(img)
                except Exception as fallback_e:
                    print(f"  Doc-Image-Tool 回退失败: {fallback_e}")
                else:
                    _write_output_image(output_path, result)
                    return output_path
            print("  回退到 OpenCV 规则管线。")
        else:
            _write_output_image(output_path, result)
            return output_path

    try:
        result = _run_opencv_mask_inpaint_pipeline(img)
    except Exception as e:
        print(f"  OpenCV mask+inpaint 管线失败: {e}，回退到旧规则管线。")
        result = _run_legacy_opencv_pipeline(img)

    _write_output_image(output_path, result)
    return output_path


def _write_output_image(output_path: Path, result: np.ndarray) -> None:
    """保存输出图，JPEG 用高质量以减少灰块/蚊噪。"""
    ext = output_path.suffix.lower()
    if ext in ('.jpg', '.jpeg'):
        cv2.imwrite(str(output_path), result, [int(cv2.IMWRITE_JPEG_QUALITY), 98])
    else:
        cv2.imwrite(str(output_path), result)


def _run_opencv_mask_inpaint_pipeline(img: np.ndarray) -> np.ndarray:
    """
    OpenCV 本地擦除管线：先生成笔迹 mask，再对 mask 区域做局部修复。

    这个后端模拟 ChatGPT/清理工具常见的处理方式，但只使用 OpenCV：
    - HSV/Lab 颜色分割找红色批改；
    - 彩色批改默认白填，避免 OpenCV inpaint 把相邻黑字扩散成黑块；
    - 横线邻域和底部草稿区用版面规则白填，避免 OpenCV inpaint 大块区域产生黑斑。
    """
    print("  阶段 1/5: 生成彩色批改 mask...")
    mask = _build_colored_ink_mask(img)
    px = int(np.count_nonzero(mask))
    total = img.shape[0] * img.shape[1]
    print(f"    彩色批改 mask: {px} px ({px / total * 100:.2f}%)")

    print("  阶段 2/5: 白填彩色批改区域（大模型/LaMa 可使用 {mask} 做生成式修复）...")
    stage1 = img.copy()
    stage1[mask > 0] = [255, 255, 255]

    print("  阶段 3/5: 背景归一化（去阴影、白化纸张）...")
    stage2 = _normalize_document_background(stage1)

    print("  阶段 4/5: 二次保守擦除横线答案与稀疏草稿...")
    stage3 = _erase_answers_around_blank_lines(stage2)
    stage3 = _erase_sparse_row_handwriting(stage3)

    print("  阶段 5/5: 保守保护题号栏并清理边缘...")
    result = stage3
    result = _restore_question_number_strip(result, _normalize_document_background(_erase_colored_ink(img)))
    return _clean_outer_photo_edges(result)


def _run_legacy_opencv_pipeline(img: np.ndarray) -> np.ndarray:
    """旧版规则后端，作为 mask+inpaint 失败时的兜底。"""
    print("  阶段 1/4: 擦除彩色笔迹（红色批改圈、标注等）...")
    try:
        stage1 = _erase_colored_ink(img)
    except Exception as e:
        print(f"  彩色笔迹擦除失败: {e}，跳过此阶段")
        stage1 = img.copy()

    print("  阶段 2/4: 背景归一化（去除光照不均、纸张发黄）...")
    try:
        stage2 = _normalize_document_background(stage1)
    except Exception as e:
        print(f"  背景归一化失败: {e}，跳过此阶段")
        stage2 = stage1

    print("  阶段 3/4: 稀疏区域黑色笔迹精细擦除...")
    try:
        result = _erase_sparse_row_handwriting(stage2)
    except Exception as e:
        print(f"  精细擦除失败: {e}，使用阶段 2 结果")
        result = stage2

    print("  阶段 4/4: 左栏二值整理（去掉灰雾与不干净笔迹残影）...")
    try:
        result = _finalize_left_strip_scan(result)
    except Exception as e:
        print(f"  左栏整理失败: {e}")
    return result


def remove_handwriting_with_model_backend(
    *,
    image_path: Path,
    img: np.ndarray,
    backend: str,
    model_repo: Path | None,
    model_checkpoint: Path | None,
    model_command: str | None,
    model_device: str,
    mask_threshold: float,
) -> np.ndarray:
    """
    模型后端统一入口。

    支持四种接入方式：
    1. `doc-image-tool`：使用已接入的 Sauvola + K-Means 文档清理模型/工具链；
    2. `--model-command`：最稳妥，适配 DeepLabV3+/DIS/EraseNet/LaMa 各仓库自己的预测脚本；
       命令模板中可用 `{input}` `{staged_input}` `{mask}` `{output}` `{input_dir}` `{output_dir}`
       `{repo}` `{checkpoint}` `{device}`。
    3. `lama`：未提供 `--model-command` 时，自动尝试 IOPaint/lama-cleaner 的 LaMa inpainting。
    4. `--model-checkpoint`：当权重是 TorchScript 或 `torch.save(model)` 的完整模型时，
       直接前向预测手写 mask，再做白填。

    不直接假设第三方仓库内部类名，因为这三套实现的工程结构、权重格式都不同。
    """
    backend = backend.lower().replace('_', '-')
    if backend not in {'doc-image-tool', 'deeplabv3plus', 'dis', 'erasenet', 'lama', 'torchscript'}:
        raise ValueError(f"未知模型后端: {backend}")

    if backend == 'doc-image-tool':
        return _run_doc_image_tool_pipeline(img)

    if backend == 'lama' and not model_command:
        return _run_lama_inpaint_backend(
            image_path=image_path,
            img=img,
            model_device=model_device,
        )

    if (
        backend == 'deeplabv3plus'
        and model_repo
        and model_checkpoint
        and _should_use_deeplab_repo_adapter(model_repo, model_command)
    ):
        return _run_deeplabv3plus_repo_model(
            image_path=image_path,
            model_repo=model_repo,
            model_checkpoint=model_checkpoint,
            model_device=model_device,
        )

    if model_command:
        return _run_external_model_command(
            image_path=image_path,
            img_shape=img.shape,
            backend=backend,
            model_repo=model_repo,
            model_checkpoint=model_checkpoint,
            model_command=model_command,
            model_device=model_device,
        )

    if model_checkpoint:
        return _run_torch_mask_model(
            img=img,
            checkpoint=model_checkpoint,
            device=model_device,
            mask_threshold=mask_threshold,
        )

    raise RuntimeError(
        f"{backend} 后端需要提供 --model-command 或 --model-checkpoint。"
        "这三个开源项目的权重/入口不统一，建议先用仓库自带预测脚本拼成命令模板。"
    )


def _should_use_deeplab_repo_adapter(model_repo: Path, model_command: str | None) -> bool:
    """HandWritingEraser-Pytorch 没有 infer.py 时，走仓库自带 predict_one.py。"""
    repo = Path(model_repo).expanduser().resolve()
    if not (repo / 'predict_one.py').exists():
        return False
    if not model_command:
        return True
    return 'infer.py' in model_command and not (repo / 'infer.py').exists()


def _run_deeplabv3plus_repo_model(
    *,
    image_path: Path,
    model_repo: Path,
    model_checkpoint: Path,
    model_device: str,
) -> np.ndarray:
    """适配 AndSonder/HandWritingEraser-Pytorch 的 predict_one.py 单图推理。"""
    repo = Path(model_repo).expanduser().resolve()
    checkpoint = Path(model_checkpoint).expanduser().resolve()
    predict_one = repo / 'predict_one.py'
    if not predict_one.exists():
        raise FileNotFoundError(f"Deeplab 仓库缺少 predict_one.py: {predict_one}")
    if not checkpoint.exists():
        raise FileNotFoundError(f"模型权重不存在: {checkpoint}")

    try:
        import torch
        import torch.nn as nn
        from torchvision import transforms as T
    except ImportError as e:
        raise RuntimeError("DeeplabV3+ 后端需要安装 torch 和 torchvision") from e

    device_text = str(model_device)
    if device_text == 'auto':
        dev = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    elif device_text.isdigit():
        os.environ['CUDA_VISIBLE_DEVICES'] = device_text
        dev = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    else:
        dev = torch.device(device_text)

    print(f"  HandWritingEraser-Pytorch 无 infer.py，改用 predict_one.py 原生推理: {dev}")
    sys.path.insert(0, str(repo))
    old_cwd = os.getcwd()
    try:
        spec = importlib.util.spec_from_file_location('_hwe_predict_one', predict_one)
        if spec is None or spec.loader is None:
            raise RuntimeError(f"无法加载 Deeplab 推理脚本: {predict_one}")
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)

        transform = T.Compose([
            T.ToTensor(),
            T.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]),
        ])
        model = module.network.modeling.deeplabv3plus_resnet101(num_classes=3, output_stride=16)
        loaded = torch.load(str(checkpoint), map_location=torch.device('cpu'))
        state = loaded.get('model_state', loaded) if isinstance(loaded, dict) else loaded
        model.load_state_dict(state)
        model = nn.DataParallel(model)
        model.to(dev)
        model.eval()

        opts = argparse.Namespace(device=dev, transform=transform)
        with tempfile.TemporaryDirectory(prefix='deeplabv3plus_predict_') as tmp:
            os.chdir(tmp)
            rgb = module.erase_hand_write(str(Path(image_path).expanduser().resolve()), model, opts)
    finally:
        os.chdir(old_cwd)
        try:
            sys.path.remove(str(repo))
        except ValueError:
            pass

    if rgb is None:
        raise RuntimeError("DeeplabV3+ 推理未返回图片")
    return cv2.cvtColor(rgb.astype(np.uint8), cv2.COLOR_RGB2BGR)


def _run_doc_image_tool_pipeline(img: np.ndarray) -> np.ndarray:
    """使用 Doc-Image-Tool 的文档漂白和笔记去噪，再执行本地精细擦除。"""
    print("  阶段 1/6: 擦除彩色笔迹（红色批改圈、标注等）...")
    stage1 = _erase_colored_ink(img)
    question_strip_source = _normalize_document_background(stage1)

    print("  阶段 2/6: Doc-Image-Tool Sauvola 文档漂白...")
    binary_bleach = sauvola_threshold(stage1, window_size=15, k=0.2, r=128)
    stage2 = np.full_like(stage1, 255)
    text_mask = binary_bleach == 0
    stage2[text_mask] = stage1[text_mask]

    print("  阶段 3/6: Doc-Image-Tool K-Means 笔记去噪美化...")
    opts = _DenoiseOptions(
        num_colors=6,
        white_bg=True,
        saturate=True,
        quiet=True,
        sample_fraction=0.05,
        value_threshold=0.25,
        sat_threshold=0.20,
    )
    stage3 = docscan_main(stage2, opts)

    print("  阶段 4/6: 填空横线附近答案擦除...")
    stage4 = _erase_answers_around_blank_lines(stage3)

    print("  阶段 5/6: 保守擦除行尾填空与底部草稿...")
    result = _erase_sparse_row_handwriting(stage4, erase_sparse_rows=False)

    print("  阶段 6/6: 保守增强印刷文字并清理照片边缘...")
    result = _boost_print_contrast(result)
    result = _restore_question_number_strip(result, question_strip_source)
    return _clean_outer_photo_edges(result)


def _run_lama_inpaint_backend(
    *,
    image_path: Path,
    img: np.ndarray,
    model_device: str,
) -> np.ndarray:
    """
    使用 LaMa 做生成式局部修复。

    默认优先调用 IOPaint：
      pip install iopaint
      python exam_scanner.py ./imgs --erase --erase-backend lama

    若本机仍在使用旧版 lama-cleaner，也会自动尝试 `lama-cleaner` 命令。
    """
    source_image = Path(image_path).expanduser().resolve()
    with tempfile.TemporaryDirectory(prefix='lama_inpaint_') as tmp:
        tmp_dir = Path(tmp)
        mask_path = tmp_dir / 'mask.png'
        output_path = tmp_dir / 'lama_output.png'
        mask = _build_handwriting_erase_mask(img)
        cv2.imwrite(str(mask_path), mask)

        if importlib.util.find_spec('simple_lama_inpainting') is not None:
            print("  使用 simple-lama-inpainting 直接执行 LaMa 修复...")
            out = _run_simple_lama_inpainting(img, mask, model_device)
            return _postprocess_lama_document_result(out, img)

        commands = _candidate_lama_commands(
            image_path=source_image,
            mask_path=mask_path,
            output_path=output_path,
            model_device=model_device,
        )
        if not commands:
            raise RuntimeError(
                "未找到 LaMa 工具。请先安装 IOPaint: pip install iopaint，"
                "或使用 --model-command 显式传入 LaMa 推理命令。"
            )

        errors = []
        for command in commands:
            print("  调用 LaMa 命令: " + " ".join(str(part) for part in command))
            proc = subprocess.run(
                [str(part) for part in command],
                text=True,
                capture_output=True,
            )
            if proc.stdout.strip():
                print(proc.stdout.strip())
            if proc.returncode == 0 and output_path.exists():
                out = cv2.imread(str(output_path), cv2.IMREAD_UNCHANGED)
                if out is None:
                    errors.append(f"命令成功但无法读取输出: {output_path}")
                    continue
                if out.ndim == 2:
                    out = cv2.cvtColor(out, cv2.COLOR_GRAY2BGR)
                elif out.shape[2] == 4:
                    out = cv2.cvtColor(out, cv2.COLOR_BGRA2BGR)
                if out.shape[:2] != img.shape[:2]:
                    out = cv2.resize(out, (img.shape[1], img.shape[0]), interpolation=cv2.INTER_AREA)
                return _postprocess_lama_document_result(out, img)

            errors.append(proc.stderr.strip() or f"退出码 {proc.returncode}")

        raise RuntimeError("LaMa 调用失败: " + " | ".join(errors))


def _run_simple_lama_inpainting(
    img: np.ndarray,
    mask: np.ndarray,
    model_device: str,
) -> np.ndarray:
    """直接调用 simple-lama-inpainting 包，避免依赖 IOPaint CLI。"""
    global _simple_lama_instance
    try:
        import torch
        from simple_lama_inpainting import SimpleLama
    except ImportError as e:
        raise RuntimeError("simple-lama-inpainting 后端需要安装 torch 和 simple-lama-inpainting") from e

    if model_device == 'auto':
        if torch.cuda.is_available():
            device = torch.device('cuda')
        elif hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
            device = torch.device('mps')
        else:
            device = torch.device('cpu')
    else:
        device = torch.device(model_device)

    if _simple_lama_instance is None:
        print(f"  加载 LaMa 模型: {device}")
        _simple_lama_instance = SimpleLama(device=device)

    rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    mask_u8 = (mask > 0).astype(np.uint8) * 255
    result = _simple_lama_instance(PILImage.fromarray(rgb), PILImage.fromarray(mask_u8))
    result_rgb = np.asarray(result).astype(np.uint8)
    if result_rgb.ndim == 2:
        return cv2.cvtColor(result_rgb, cv2.COLOR_GRAY2BGR)
    return cv2.cvtColor(result_rgb, cv2.COLOR_RGB2BGR)


def _candidate_lama_commands(
    *,
    image_path: Path,
    mask_path: Path,
    output_path: Path,
    model_device: str,
) -> list[list[str | Path]]:
    """按当前环境返回可尝试的 LaMa CLI 命令。"""
    device = 'cpu' if model_device == 'auto' else model_device
    common_args = [
        '--model=lama',
        f'--device={device}',
        f'--image={image_path}',
        f'--mask={mask_path}',
        f'--output={output_path}',
    ]
    commands: list[list[str | Path]] = []

    if shutil.which('iopaint'):
        commands.append(['iopaint', 'run', *common_args])
    if importlib.util.find_spec('iopaint') is not None:
        commands.append([sys.executable, '-m', 'iopaint', 'run', *common_args])

    # lama-cleaner 是 IOPaint 的旧包名，保留兼容。
    if shutil.which('lama-cleaner'):
        commands.append(['lama-cleaner', *common_args])
    if importlib.util.find_spec('lama_cleaner') is not None:
        commands.append([sys.executable, '-m', 'lama_cleaner', *common_args])

    return commands


def _postprocess_lama_document_result(result: np.ndarray, original: np.ndarray) -> np.ndarray:
    """LaMa 输出后做轻量文档白化，并恢复左侧题号栏，避免生成式模型改写题号。"""
    normalized = _normalize_document_background(result)
    question_strip_source = _normalize_document_background(_erase_colored_ink(original))
    normalized = _restore_question_number_strip(normalized, question_strip_source)
    return _clean_outer_photo_edges(normalized)


def _run_external_model_command(
    *,
    image_path: Path,
    img_shape: tuple[int, ...],
    backend: str,
    model_repo: Path | None,
    model_checkpoint: Path | None,
    model_command: str,
    model_device: str,
) -> np.ndarray:
    """
    调用第三方仓库自带预测脚本。

    例：
      --erase-backend deeplabv3plus \
      --model-repo third-party-libs/HandWritingEraser-Pytorch \
      --model-checkpoint checkpoints/best.pth \
      --model-command 'python {repo}/infer.py --input {input} --mask {mask} --output {output} --ckpt {checkpoint}'

    对只支持批量目录输入的仓库，可用 `{input_dir}` 和 `{output_dir}`；
    LaMa/扩散修复类仓库可直接使用 `{mask}` 作为修复 mask。
    """
    repo = Path(model_repo).expanduser().resolve() if model_repo else None
    checkpoint = Path(model_checkpoint).expanduser().resolve() if model_checkpoint else None
    source_image = Path(image_path).expanduser().resolve()

    with tempfile.TemporaryDirectory(prefix=f'{backend}_erase_') as tmp:
        tmp_dir = Path(tmp)
        output_path = tmp_dir / 'model_output.png'
        input_dir = tmp_dir / 'input'
        output_dir = tmp_dir / 'output'
        input_dir.mkdir()
        output_dir.mkdir()
        staged_input = input_dir / source_image.name
        shutil.copy2(source_image, staged_input)
        original = cv2.imread(str(source_image))
        if original is None:
            raise RuntimeError(f"无法读取输入图片以生成 mask: {source_image}")
        mask_path = tmp_dir / 'mask.png'
        cv2.imwrite(str(mask_path), _build_handwriting_erase_mask(original))
        fmt = {
            'input': str(source_image),
            'staged_input': str(staged_input),
            'mask': str(mask_path),
            'output': str(output_path),
            'input_dir': str(input_dir),
            'output_dir': str(output_dir),
            'repo': str(repo) if repo else '',
            'checkpoint': str(checkpoint) if checkpoint else '',
            'device': model_device,
        }
        command = model_command.format(**fmt)
        print(f"  调用模型命令: {command}")
        proc = subprocess.run(
            command,
            shell=True,
            cwd=str(repo) if repo else None,
            text=True,
            capture_output=True,
        )
        if proc.stdout.strip():
            print(proc.stdout.strip())
        if proc.returncode != 0:
            raise RuntimeError(proc.stderr.strip() or f"模型命令退出码 {proc.returncode}")
        if not output_path.exists():
            candidates = sorted([
                p for p in output_dir.iterdir()
                if p.suffix.lower() in {'.png', '.jpg', '.jpeg', '.bmp', '.webp', '.tiff'}
            ])
            if candidates:
                output_path = candidates[0]
            else:
                raise RuntimeError(
                    f"模型命令未生成输出文件: {output_path}，输出目录也为空: {output_dir}"
                )

        out = cv2.imread(str(output_path), cv2.IMREAD_UNCHANGED)
        if out is None:
            raise RuntimeError(f"无法读取模型输出: {output_path}")

        if out.ndim == 2:
            # 若脚本输出的是 mask，则用白填方式擦除。
            return _apply_handwriting_mask(original, out)

        if out.shape[:2] != img_shape[:2]:
            out = cv2.resize(out, (img_shape[1], img_shape[0]), interpolation=cv2.INTER_AREA)
        if out.shape[2] == 4:
            out = cv2.cvtColor(out, cv2.COLOR_BGRA2BGR)
        if backend == 'lama':
            return _postprocess_lama_document_result(out, original)
        return out


def _run_torch_mask_model(
    *,
    img: np.ndarray,
    checkpoint: Path,
    device: str,
    mask_threshold: float,
) -> np.ndarray:
    """
    直接运行 TorchScript / 完整 PyTorch 模型，约定输出为手写区域 mask。

    适合把 DeepLabV3+、DIS、EraseNet 导出成 TorchScript 后接入。
    普通 state_dict 需要原仓库模型类定义，此函数不会猜测架构。
    """
    try:
        import torch
    except ImportError as e:
        raise RuntimeError("模型后端需要安装 torch") from e

    checkpoint = Path(checkpoint)
    if not checkpoint.exists():
        raise FileNotFoundError(f"模型权重不存在: {checkpoint}")

    dev = torch.device(device if device != 'auto' else ('cuda' if torch.cuda.is_available() else 'cpu'))

    try:
        model = torch.jit.load(str(checkpoint), map_location=dev)
    except Exception:
        loaded = torch.load(str(checkpoint), map_location=dev)
        if isinstance(loaded, torch.nn.Module):
            model = loaded
        else:
            raise RuntimeError(
                "该 checkpoint 看起来是 state_dict，无法脱离原仓库模型类直接推理。"
                "请改用 --model-command 调原仓库预测脚本，或导出 TorchScript。"
            )

    model.eval()
    rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    tensor = torch.from_numpy(rgb).float().permute(2, 0, 1).unsqueeze(0) / 255.0
    tensor = tensor.to(dev)

    with torch.no_grad():
        pred = model(tensor)
        if isinstance(pred, (list, tuple)):
            pred = pred[0]
        if isinstance(pred, dict):
            pred = next(iter(pred.values()))
        if pred.ndim == 4:
            pred = pred[:, :1, :, :]
        pred = torch.sigmoid(pred).squeeze().detach().cpu().numpy()

    pred = cv2.resize(pred.astype(np.float32), (img.shape[1], img.shape[0]))
    mask = (pred >= mask_threshold).astype(np.uint8) * 255
    return _apply_handwriting_mask(img, mask)


def _apply_handwriting_mask(img: np.ndarray, mask: np.ndarray) -> np.ndarray:
    """模型输出 handwriting mask 后的统一擦除策略。"""
    if mask.ndim == 3:
        mask = cv2.cvtColor(mask, cv2.COLOR_BGR2GRAY)
    mask = cv2.resize(mask, (img.shape[1], img.shape[0]), interpolation=cv2.INTER_NEAREST)
    _, mask = cv2.threshold(mask, 127, 255, cv2.THRESH_BINARY)
    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN, k, iterations=1)
    mask = cv2.dilate(mask, k, iterations=1)

    # 文档去手写通常白填比自然图 inpaint 更干净；之后再做一次背景归一化。
    result = img.copy()
    result[mask > 0] = [255, 255, 255]
    return _normalize_document_background(result)


def _build_handwriting_erase_mask(img: np.ndarray) -> np.ndarray:
    """合成需要擦除的像素 mask，供 OpenCV inpaint 或外部 LaMa 类模型使用。"""
    h, w = img.shape[:2]
    mask = np.zeros((h, w), dtype=np.uint8)

    colored = _build_colored_ink_mask(img)
    mask = cv2.bitwise_or(mask, colored)

    normalized = _normalize_document_background(img)
    blank_answers = _build_blank_line_answer_mask(normalized, verbose=False)
    bottom_scratch = _build_bottom_solution_mask(normalized)
    mask = cv2.bitwise_or(mask, blank_answers)
    mask = cv2.bitwise_or(mask, bottom_scratch)

    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, k, iterations=1)
    mask = cv2.dilate(mask, k, iterations=1)

    px = int(np.count_nonzero(mask))
    total = h * w
    print(f"    合成擦除 mask: {px} px ({px / total * 100:.2f}%)")
    return mask


def _build_colored_ink_mask(img: np.ndarray) -> np.ndarray:
    """
    检测高饱和批改笔迹 mask。

    以教师红笔为主，略带蓝/紫笔容错；只做像素候选，后续由版面阶段继续保护印刷内容。
    """
    hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)

    red_lo1 = np.array([0, 100, 70], dtype=np.uint8)
    red_hi1 = np.array([7, 255, 255], dtype=np.uint8)
    red_lo2 = np.array([168, 100, 70], dtype=np.uint8)
    red_hi2 = np.array([180, 255, 255], dtype=np.uint8)
    red = cv2.bitwise_or(
        cv2.inRange(hsv, red_lo1, red_hi1),
        cv2.inRange(hsv, red_lo2, red_hi2),
    )

    # 少量容错：蓝紫色笔迹常见于学生改答案，但阈值收紧，避免误伤灰色/黑色印刷。
    blue_purple = cv2.inRange(
        hsv,
        np.array([105, 80, 45], dtype=np.uint8),
        np.array([155, 255, 230], dtype=np.uint8),
    )

    mask = cv2.bitwise_or(red, blue_purple)
    mask = _filter_small_pen_components(mask)
    return mask


def _filter_small_pen_components(mask: np.ndarray) -> np.ndarray:
    """过滤颜色分割噪点和大面积印刷色块，只保留更像笔画的连通域。"""
    num, labels, stats, _ = cv2.connectedComponentsWithStats(mask, connectivity=8)
    filtered = np.zeros_like(mask)
    for i in range(1, num):
        x, y, ww, hh, area = stats[i]
        if area < 8:
            continue
        fill_ratio = area / float(max(1, ww * hh))
        long_side = max(ww, hh)
        short_side = min(ww, hh)
        looks_like_stroke = (
            area <= 3500
            and long_side >= 3
            and short_side <= 90
            and fill_ratio <= 0.72
        )
        if looks_like_stroke:
            filtered[labels == i] = 255
    return filtered


def _inpaint_document_mask(img: np.ndarray, mask: np.ndarray) -> np.ndarray:
    """对 mask 区域做文档场景的局部修复；无 mask 时直接返回副本。"""
    if mask.ndim == 3:
        mask = cv2.cvtColor(mask, cv2.COLOR_BGR2GRAY)
    _, mask = cv2.threshold(mask, 127, 255, cv2.THRESH_BINARY)
    px = int(np.count_nonzero(mask))
    if px < 50:
        print("    mask 太小，跳过 inpaint")
        return img.copy()

    radius = 3 if px < img.shape[0] * img.shape[1] * 0.08 else 5
    repaired = cv2.inpaint(img, mask, radius, cv2.INPAINT_TELEA)

    # 大块底部草稿区在文档里应接近空白，inpaint 后再白填能避免纸纹被拖出脏影。
    large_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (33, 33))
    large_regions = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, large_kernel, iterations=2)
    repaired[large_regions > 0] = np.maximum(repaired[large_regions > 0], 245)
    return repaired


def _erase_colored_ink(img: np.ndarray) -> np.ndarray:
    """
    用 HSV 颜色空间检测并白填擦除手写彩色墨水（教师批改红笔）。

    只针对「高饱和度鲜红色」（HSV H:0-7° / 168-180°, S≥100, V≥70）：
    - 捕获：教师用红笔圈出的错题圆圈、叉号、旁注「错」字
    - 跳过：试卷印刷设计中的珊瑚橙色题号圆圈（S 较低 / 色相偏橙）

    用白色直接填充（不用 inpaint），避免颜色重建产生灰色污点。

    Returns:
        擦除彩色笔迹后的图像
    """
    colored_mask = _build_colored_ink_mask(img)

    px = int(np.count_nonzero(colored_mask))
    total = img.shape[0] * img.shape[1]
    print(f"    检测到红色笔迹: {px} px ({px / total * 100:.2f}%)")

    if px < 50:
        return img.copy()

    # 小核少量膨胀：只盖住红笔画本身，避免膨胀进左侧印刷橙色题号圆
    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    colored_mask = cv2.dilate(colored_mask, k, iterations=1)

    # 白色直接填充（不用 inpaint）：避免灰色重建污点
    result = img.copy()
    result[colored_mask > 0] = [255, 255, 255]
    return result


def _normalize_document_background(img: np.ndarray) -> np.ndarray:
    """
    背景归一化：让纸张趋近纯白、印刷墨迹趋近黑色。

    重要：不得在「已经把背景设为 255」之后再做全局锐化 —— Laplacian/叠加会把
    白色边缘拉回 180~240 的灰值，整页出现「灰色污染」（用户反馈的主要问题）。
    """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    bg_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (61, 61))
    background = cv2.dilate(gray, bg_kernel)
    background = np.where(background == 0, 1, background).astype(np.float32)

    norm = gray.astype(np.float32) / background * 255.0
    norm = np.clip(norm, 0, 255).astype(np.uint8)

    # 背景硬白化（略放宽，压住照片阴影）
    norm[norm > 210] = 255

    result_gray = _squeeze_gray_to_bnw(norm)

    return cv2.cvtColor(result_gray, cv2.COLOR_GRAY2BGR)


def _finalize_left_strip_scan(
    img: np.ndarray,
    *,
    band_frac: float = 0.72,
) -> np.ndarray:
    """
    左栏题干带：除「疑似印刷墨迹」外一律 #FFF，避免 γ/半调产生灰块与笔迹残影。

    墨迹 = 自适应二值 ∪ (灰度 < 低阈) 做小开运算去椒盐；墨迹像素压到近黑，其余白。
    右栏 (>band_frac) 原样保留。
    """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    cut = max(24, min(w - 24, int(round(w * band_frac))))
    lc = gray[:, :cut].copy()
    right = gray[:, cut:]

    blk = 33
    atk = cv2.adaptiveThreshold(
        lc,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY_INV,
        blk,
        11,
    )
    ink = (atk > 0) | (lc < 148)
    ink_u8 = (ink.astype(np.uint8) * 255)
    k2 = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
    ink_u8 = cv2.morphologyEx(ink_u8, cv2.MORPH_OPEN, k2)
    ink = ink_u8 > 0

    new_left = np.full_like(lc, 255, dtype=np.uint8)
    v = lc.astype(np.int32)
    new_left[ink] = np.clip((v[ink] - 25) * 18 // 25, 0, 95).astype(np.uint8)

    full = np.hstack([new_left, right])
    return cv2.cvtColor(full, cv2.COLOR_GRAY2BGR)


def _erase_answers_around_blank_lines(
    img: np.ndarray,
    *,
    band_frac: float = 0.74,
) -> np.ndarray:
    """
    对填空横线附近做局部擦除：保留长横线，擦掉横线上下的短手写笔画。

    模型/漂白后，填空答案通常仍以短促黑色连通域压在下划线上；
    这里先用横向形态学找出答题线，再只在答题线的窄带内白填非横线墨迹。
    """
    erase_full = _build_blank_line_answer_mask(img, band_frac=band_frac, verbose=True)
    if np.count_nonzero(erase_full) == 0:
        return img

    result = img.copy()
    result[erase_full > 0] = [255, 255, 255]
    return result


def _build_blank_line_answer_mask(
    img: np.ndarray,
    *,
    band_frac: float = 0.74,
    verbose: bool = True,
) -> np.ndarray:
    """识别填空横线附近的短手写答案，返回待擦除 mask。"""
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    cut = max(24, min(w - 1, int(round(w * band_frac))))
    left = gray[:, :cut]

    _, dark = cv2.threshold(left, 178, 255, cv2.THRESH_BINARY_INV)
    line_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (35, 1))
    hlines = cv2.morphologyEx(dark, cv2.MORPH_OPEN, line_kernel, iterations=1)

    # 过滤掉短噪声，只保留真实填空线/题中横线。
    num, labels, stats, _ = cv2.connectedComponentsWithStats(hlines, connectivity=8)
    keep_lines = np.zeros_like(hlines)
    for i in range(1, num):
        x, y, ww, hh, area = stats[i]
        if ww >= 28 and hh <= 6 and area >= 20:
            keep_lines[labels == i] = 255

    if np.count_nonzero(keep_lines) == 0:
        if verbose:
            print("    步骤A0 未检测到可处理填空线")
        return np.zeros((h, w), dtype=np.uint8)

    band_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (19, 21))
    line_band = cv2.dilate(keep_lines, band_kernel, iterations=1)
    protected_line = cv2.dilate(
        keep_lines,
        cv2.getStructuringElement(cv2.MORPH_RECT, (7, 3)),
        iterations=1,
    )

    erase = cv2.bitwise_and(dark, line_band)
    erase[protected_line > 0] = 0

    # 只擦横线附近较小的连通域；大块/长条多半是题干、图形或横线残段。
    num_e, labels_e, stats_e, _ = cv2.connectedComponentsWithStats(erase, connectivity=8)
    filtered = np.zeros_like(erase)
    for i in range(1, num_e):
        x, y, ww, hh, area = stats_e[i]
        if 3 <= area <= 520 and ww <= 95 and hh <= 28:
            filtered[labels_e == i] = 255
    erase = filtered
    erase = cv2.dilate(
        erase,
        cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3)),
        iterations=1,
    )

    erased_px = int(np.count_nonzero(erase))
    if verbose:
        print(f"    步骤A0 填空线邻域白填像素: {erased_px} px")
    erase_full = np.zeros((h, w), dtype=np.uint8)
    erase_full[:, :cut] = erase
    return erase_full


def _clean_outer_photo_edges(img: np.ndarray) -> np.ndarray:
    """清掉手机拍摄边缘的黑边/侧边噪声，避免影响与扫描效果图的观感。"""
    result = img.copy()
    h, w = result.shape[:2]
    # 左侧题号很靠边，不能整条大面积白填；只清最外缘拍摄黑边。
    edge_x = max(24, int(round(w * 0.012)))
    edge_y = max(8, int(round(h * 0.006)))
    result[:, :edge_x] = [255, 255, 255]
    result[:, w - max(8, edge_x // 2):] = [255, 255, 255]
    result[:edge_y, :] = [255, 255, 255]
    result[h - edge_y:, :] = [255, 255, 255]
    return result


def _restore_question_number_strip(
    processed: np.ndarray,
    source: np.ndarray,
    *,
    strip_frac: float = 0.058,
) -> np.ndarray:
    """恢复左侧题号栏，避免灰色圆圈题号被模型/阈值化当成背景抹掉。"""
    result = processed.copy()
    h, w = result.shape[:2]
    strip_w = max(84, min(int(round(w * strip_frac)), 118, w // 5))
    source_resized = source
    if source_resized.shape[:2] != result.shape[:2]:
        source_resized = cv2.resize(source_resized, (w, h), interpolation=cv2.INTER_AREA)

    result[:, :strip_w] = source_resized[:, :strip_w]
    print(f"    题号栏保护: 恢复左侧 {strip_w}px")
    return result


def _boost_print_contrast(img: np.ndarray) -> np.ndarray:
    """只做灰度拉伸，不做开运算删点，避免题干细笔画缺失。"""
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    out = gray.astype(np.float32)
    out = (out - 135.0) / (235.0 - 135.0) * 255.0
    out = np.clip(out, 0, 255).astype(np.uint8)
    out[out > 232] = 255
    return cv2.cvtColor(out, cv2.COLOR_GRAY2BGR)


def _squeeze_gray_to_bnw(gray: np.ndarray) -> np.ndarray:
    """
    将「接近纸白」的中间灰.squeeze 成 255，将「墨迹」.squeeze 得更黑。
    减少整页雾霾感，且不引入锐化灰边。
    """
    lo, hi = 115, 238
    out = gray.astype(np.float32)
    out = (out.astype(np.float32) - lo) / (hi - lo) * 255.0
    out = np.clip(out, 0, 255).astype(np.uint8)
    out[out > 200] = 255
    return out


def _erase_sparse_row_handwriting(
    img: np.ndarray,
    *,
    erase_sparse_rows: bool = True,
) -> np.ndarray:
    """
    左栏填空与笔迹擦除：

    - 全程仅在「题干投影宽」proj_w (~72%) 内统计密度，图示列不参与；
    - 整行稀疏行：与印刷密行邻接保护带外，低阈二值化后白填；
    - 「左半题干密 + 右半行尾疏」分行：同一行上题干与手写答案分列，避免保护带挡掉填空笔迹。
    """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape

    proj_w = max(24, min(w - 1, int(round(w * 0.72))))
    gp = gray[:, :proj_w]

    _, bin_for_rows = cv2.threshold(gp, 198, 255, cv2.THRESH_BINARY_INV)
    row_black = bin_for_rows.sum(axis=1).astype(np.float32) / 255.0
    row_density = row_black / proj_w

    PRINT_THRESH = 0.033
    SPARSE_LOW = 0.0015
    SPARSE_HIGH = PRINT_THRESH

    is_sparse = (row_density > SPARSE_LOW) & (row_density < SPARSE_HIGH)

    is_dense = row_density >= PRINT_THRESH
    protect = np.zeros(h, dtype=bool)
    for y in range(h):
        if is_dense[y]:
            protect[max(0, y - 3):min(h, y + 4)] = True
    is_sparse = is_sparse & ~protect
    if not erase_sparse_rows:
        is_sparse[:] = False

    sparse_count = int(is_sparse.sum())
    mode = "启用" if erase_sparse_rows else "禁用"
    print(f"    步骤A 整行稀疏: {sparse_count} 行（{mode}）")

    split_c = max(100, min(proj_w - 80, int(round(proj_w * 0.48))))
    _, bl_hs = cv2.threshold(gp[:, :split_c], 198, 255, cv2.THRESH_BINARY_INV)
    _, br_hs = cv2.threshold(gp[:, split_c:], 200, 255, cv2.THRESH_BINARY_INV)
    rw = float(split_c)
    rrest = float(proj_w - split_c)
    d_left = bl_hs.sum(axis=1).astype(np.float32) / (255.0 * rw)
    d_right = br_hs.sum(axis=1).astype(np.float32) / (255.0 * max(1.0, rrest))

    row_idx = np.arange(h, dtype=np.int32)

    # 行尾填空区：左侧印刷密、右侧仅少量墨迹；**不用**整条「密行保护带」封杀，
    # 只改写右半条，题干汉字不会被波及。
    is_partial = (
        (row_idx > int(h * 0.05))
        & (d_left >= 0.028)
        & (d_right > 0.0025)
        & (d_right < 0.055)
    )
    partial_n = int(is_partial.sum())
    print(f"    步骤A 行尾填空分区: {partial_n} 行")

    erase_x_end = proj_w

    sparse_2d = np.zeros((h, w), dtype=np.uint8)
    sparse_2d[is_sparse, :erase_x_end] = 255

    partial_2d = np.zeros((h, w), dtype=np.uint8)
    partial_2d[is_partial, split_c:erase_x_end] = 255

    combined = cv2.bitwise_or(sparse_2d[:, :erase_x_end], partial_2d[:, :erase_x_end])
    _, bin_erase = cv2.threshold(gray[:, :erase_x_end], 110, 255, cv2.THRESH_BINARY_INV)

    erase_mask_left = cv2.bitwise_and(bin_erase, combined)
    k = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
    erase_mask_left = cv2.dilate(erase_mask_left, k, iterations=1)

    erased_px = int(np.count_nonzero(erase_mask_left))
    print(f"    步骤A 白填像素（左栏）: {erased_px} px")

    result = img.copy()
    erase_full = np.zeros((h, w), dtype=np.uint8)
    erase_full[:, :erase_x_end] = erase_mask_left
    result[erase_full > 0] = [255, 255, 255]

    _, bin_sol = cv2.threshold(gp, 203, 255, cv2.THRESH_BINARY_INV)
    row_density_sol = bin_sol.sum(axis=1).astype(np.float32) / (255.0 * proj_w)

    result = _erase_solution_block(result, row_density_sol, h, w)
    return result


def _build_bottom_solution_mask(img: np.ndarray) -> np.ndarray:
    """识别页面下半部的大块手写草稿区，返回整块擦除 mask。"""
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    proj_w = max(24, min(w - 1, int(round(w * 0.72))))
    gp = gray[:, :proj_w]
    _, bin_sol = cv2.threshold(gp, 203, 255, cv2.THRESH_BINARY_INV)
    row_density = bin_sol.sum(axis=1).astype(np.float32) / (255.0 * proj_w)

    search_top = int(h * 0.62)
    window = 30
    handwrite_lo = 0.0005
    handwrite_hi = 0.08
    min_rows = 120
    mask = np.zeros((h, w), dtype=np.uint8)

    if h - window <= search_top:
        return mask

    run_start = None
    intervals = []
    for y in range(search_top, h - window):
        window_density = float(row_density[y : y + window].mean())
        ok = handwrite_lo < window_density < handwrite_hi
        if ok:
            if run_start is None:
                run_start = y
        elif run_start is not None:
            intervals.append((run_start, y - 1))
            run_start = None

    if run_start is not None:
        intervals.append((run_start, h - window - 1))

    long = [(a, b, b - a + 1) for a, b in intervals if (b - a + 1) >= min_rows]
    if not long:
        return mask

    a, _, _ = max(long, key=lambda t: (t[2], t[0]))
    # 生成给 LaMa/外部修复模型的 mask 时宁可保守，避免把第 21 题题干也交给模型重绘。
    solution_start = max(int(h * 0.76), int(a - max(108, window * 3)))
    dense_print_rows = int(np.count_nonzero(row_density[solution_start:] > 0.035))
    if dense_print_rows > 12:
        safe_start = max(int(h * 0.78), solution_start)
        bottom_density = float(row_density[safe_start:].mean()) if safe_start < h else 0.0
        if safe_start < h - 80 and bottom_density > 0.0025:
            mask[safe_start:, :] = 255
        return mask

    mask[solution_start:, :] = 255
    return mask


def _erase_solution_block(
    img: np.ndarray,
    row_density: np.ndarray,
    h: int,
    w: int,
) -> np.ndarray:
    """
    白填页面左栏投影识别出的「大段手写解题区」（通常位于最下方）。

    旧算法假设手写行投影密度极低，与实际公式草稿（左栏仍可较密）不符。
    现改为：在 [0.62h, h] 内找滑动窗口均值落在 (lo, hi) 内的**最长连续区间**，
    若跨度 ≥ MIN_ROWS 行则取其中最靠下的那一段的起点作为解答区起始行。

    row_density：须与 _erase_sparse_row_handwriting 一致，仅为左栏 ~72% 宽度的投影密度。
    """
    SEARCH_TOP = int(h * 0.62)
    WINDOW = 30
    HANDWRITE_LO = 0.0005
    HANDWRITE_HI = 0.08
    MIN_ROWS = 120

    result = img.copy()
    if h - WINDOW <= SEARCH_TOP:
        print("    步骤B 跳过（图像过矮）")
        return result

    run_start = None
    intervals = []

    for y in range(SEARCH_TOP, h - WINDOW):
        window_density = float(row_density[y : y + WINDOW].mean())
        ok = HANDWRITE_LO < window_density < HANDWRITE_HI
        if ok:
            if run_start is None:
                run_start = y
        else:
            if run_start is not None:
                intervals.append((run_start, y - 1))
                run_start = None

    if run_start is not None:
        intervals.append((run_start, h - WINDOW - 1))

    long = [(a, b, b - a + 1) for a, b in intervals if (b - a + 1) >= MIN_ROWS]
    if not long:
        print("    步骤B 未检测到独立解答区块")
        return result

    a, b, ln = max(long, key=lambda t: (t[2], t[0]))
    PAD_UP = max(108, WINDOW * 3)
    solution_start_raw = max(0, int(a - PAD_UP))
    solution_start = max(int(h * 0.56), solution_start_raw)
    dense_print_rows = int(np.count_nonzero(row_density[solution_start:] > 0.035))
    if dense_print_rows > 12:
        safe_start = max(int(h * 0.78), solution_start)
        bottom_density = float(row_density[safe_start:].mean()) if safe_start < h else 0.0
        if safe_start < h - 80 and bottom_density > 0.0025:
            print(
                "    步骤B 检测到底部印刷题干，改为只白填更靠下草稿区: "
                f"y≥{safe_start}（底部均值密度 {bottom_density:.4f}）"
            )
            result[safe_start:, :] = [255, 255, 255]
        else:
            print(
                "    步骤B 跳过整块白填"
                f"（y≥{solution_start} 仍有 {dense_print_rows} 行疑似印刷题干）"
            )
        return result

    print(
        f"    步骤B 解答区块: 检测窗 {a}~{b}（≈{ln} 行），上扩白填自 y≥{solution_start} 至页底"
    )

    result[solution_start:, :] = [255, 255, 255]
    return result


# ============================================================
# 文本级手写笔迹擦除（基于正则，OCR 后处理用）
# ============================================================

# 常见的手写答案提示词 - 这些后面的内容通常是手写答案
ANSWER_PREFIX_RE = re.compile(
    r'(?:答案|解答|解析|解|答|选|填|写)[:：]\s*|'  # 答案: 解答: 解: 答: 等
    r'(?:[Tt]herefore|[Bb]ecause|[Aa]nswer)[:：]\s*'  # 英文答案提示
)

# 手写填空标记：题目中的空白处，学生手写了答案
# 特征：括号/方框内的短内容，通常只有数字/字母/简单符号
HANDWRITING_FILL_RE = re.compile(
    r'(?:'
    # ( ) 或 （ ）中的手写填空（内容短，通常 ≤5 个字符，不含中文）
    r'[(（]\s*([^\u4e00-\u9fff()（）]{1,5})\s*[)）]'
    r'|'
    # 【 】中的手写填空
    r'【\s*([^\u4e00-\u9fff【】]{1,5})\s*】'
    r'|'
    # [ ] 中的手写填空
    r'\[\s*([^\u4e00-\u9fff\[\]]{1,5})\s*\]'
    r')'
)

# 明显是手写选择题答案的模式：单独成行的字母或字母组合
# 如 OCR 识别到孤立的一行 "A" "B" "C" "D" 或 "AB" "CD"
STRAY_CHOICE_RE = re.compile(
    r'^\s*[A-Da-d]{1,4}\s*$'
)

# 显著的手写标记：学生用星号、下划线等标记的内容
HANDWRITING_MARKER_RE = re.compile(
    r'\*[^*]{1,20}\*|'    # *手写标记内容*
    r'_{2,}|'             # 连续下划线（填空线）
    r'--{2,}|'            # 连续横线
    r'~~|'                # 删除线标记
    r'√|'                 # 对勾
    r'[/\\]'              # 随手画的斜线
)


def erase_handwriting(text: str) -> str:
    """
    擦除 OCR 识别到的手写答题笔迹。
    
    试卷上学生手写的答案与印刷题目混合在一起，
    OCR 无法区分。此函数通过启发式规则移除常见的手写内容：
    
    1. 填空处的手写答案（括号/方框内短内容）
    2. 答案提示词后的手写内容（如"答：xxx"）
    3. 孤立的选择题答案字母（A/B/C/D）
    4. 手写标记符号（星号、下划线、对勾等）
    
    注意：会用占位符替代被擦除的内容，保留题目结构。
    """
    if not text:
        return text
    
    # ---- 1. 移除填空括号内的手写答案 ----
    # 保留括号本身，只清除内容，避免破坏题目结构
    def _clean_fill(match):
        full = match.group(0)
        # 检查内容是否像手写答案（不含中文、不含复杂标点）
        inner = (match.group(1) or match.group(2) or match.group(3) or '').strip()
        if not inner:
            return full  # 空括号，保留
        # 如果内容包含中文，大概率是印刷内容，保留
        if re.search(r'[\u4e00-\u9fff]', inner):
            return full
        # 手写答案 — 保留括号结构，用空格替代内容
        if full.startswith('(') or full.startswith('（'):
            return '( )' if full.startswith('(') else '（ ）'
        elif full.startswith('【'):
            return '【 】'
        elif full.startswith('['):
            return '[ ]'
        return full
    
    text = HANDWRITING_FILL_RE.sub(_clean_fill, text)
    
    # ---- 2. 移除答案提示词后的手写内容 ----
    # 如 "答：3.14" → 移除 "3.14"，保留 "答：___"
    def _clean_answer_prefix(match):
        return match.group(0)  # 保留提示词本身
    
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        # 如果行以答案提示词开头，尝试截断手写部分
        m = ANSWER_PREFIX_RE.match(line)
        if m:
            prefix = m.group(0)
            rest = line[m.end():].strip()
            if rest:
                # 检查剩余部分是否像手写（短、无句号、无中文长句）
                if (len(rest) <= 15 and 
                    not re.search(r'[。！？]', rest) and
                    not re.search(r'[\u4e00-\u9fff]{4,}', rest)):
                    # 疑似手写答案，保留提示词，用占位符替代答案
                    cleaned_lines.append(prefix + '[手写答案已擦除]')
                    continue
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # ---- 3. 移除孤立的选择题答案字母 ----
    lines = text.splitlines()
    cleaned_lines = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if STRAY_CHOICE_RE.match(stripped):
            # 检查上下文：前后行是否有题目内容
            prev_is_content = (i > 0 and len(lines[i-1].strip()) > 10)
            next_is_content = (i < len(lines)-1 and len(lines[i+1].strip()) > 10)
            if prev_is_content or next_is_content:
                # 上下文有题目内容，这行很可能是手写答案 → 移除
                continue
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # ---- 4. 擦除手写标记符号 ----
    text = HANDWRITING_MARKER_RE.sub(' ', text)
    
    # ---- 5. 清理擦除后产生的多余空白 ----
    text = re.sub(r' {3,}', '  ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text

def clean_ocr_text(text: str) -> str:
    """
    清洗 OCR 识别结果，移除乱码和噪声字符。
    
    常见 OCR 乱码来源：
    - 控制字符 (U+0000-U+001F, U+007F-U+009F)
    - 零宽字符、双向控制字符
    - 私有区字符 (PUA)
    - 无法解码的替换字符 (U+FFFD)
    - 孤立的组合标记
    - OCR 产生的无意义符号序列
    """
    if not text:
        return text
    
    # 1. 移除替换字符 (U+FFFD) - 这是乱码的典型标志
    text = text.replace('\ufffd', '')
    
    # 2. 移除零宽字符和双向控制字符
    zero_width_and_bidi = set(
        chr(c) for c in range(0x200B, 0x200F + 1)  # 零宽空格等
    ) | set(
        chr(c) for c in range(0x2028, 0x202E + 1)  # 行分隔符、双向控制
    ) | set(
        chr(c) for c in range(0x2060, 0x206F + 1)  # 单词连接符、不可见字符
    ) | {
        '\ufeff',   # BOM / 零宽不换行空格
        '\ufff0',   # 其他替换字符
    }
    text = ''.join(ch for ch in text if ch not in zero_width_and_bidi)
    
    # 3. 移除其他控制字符（保留常用换行和制表符）
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # 4. 移除私有区字符 (PUA: U+E000-U+F8FF, U+F0000-U+FFFFF, U+100000-U+10FFFF)
    text = re.sub(r'[\ue000-\uf8ff\U000f0000-\U000fffff\U00100000-\U0010ffff]', '', text)
    
    # 5. 规范化 Unicode（NFC 组合形式，对中文更友好）
    text = unicodedata.normalize('NFC', text)
    
    # 6. 擦除手写答题笔迹（填空答案、选择项、答案提示后的手写内容等）
    text = erase_handwriting(text)
    
    # 7. 清理 OCR 常见噪声：连续重复标点、无意义符号序列
    text = re.sub(r'[。，、；：？！…]{4,}', '...', text)  # 过多标点合并
    # 注意：不删除孤立的 1-2 个英文字母，数学公式变量（x, y, r, n, π 等）会被破坏
    text = re.sub(r'[^\S\n]{3,}', '  ', text)  # 多个空格合并
    
    # 8. 清理 OCR 中常见的错误识别模式
    # 中文被识别为乱码拼音组合
    text = re.sub(r'[a-z]{15,}', '', text, flags=re.IGNORECASE)  # 超长无意义字母串
    
    # 9. 移除垃圾行（伪中文乱码、无意义符号行、OCR 幻想文本）
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append('')
            continue
        
        ratio = _line_valid_char_ratio(stripped)
        
        # 规则 A: 有效字符比例太低 → 垃圾
        if len(stripped) >= 5 and ratio < 0.35:
            continue
        
        # 规则 B: 短行（5-25 字符）全是符号/数字/乱码 → 垃圾
        if 5 <= len(stripped) <= 25 and ratio < 0.5:
            continue
        
        # 规则 C: 包含大量 CJK 字符但不通顺 → 伪中文乱码
        cjk_count = sum(1 for ch in stripped if '\u4e00' <= ch <= '\u9fff')
        if cjk_count >= 10 and not is_coherent_chinese(stripped):
            continue
        
        # 规则 D: 超长无意义行
        if len(stripped) > 80 and ratio < 0.6:
            continue
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def is_coherent_chinese(text: str) -> bool:
    """
    判断文本是否为通顺的中文（而非 OCR 产生的伪中文乱码）。
    
    原理：
    - 真实中文高频重用常见字（的、是、一、不、了...），
      字符重复率较高。
    - OCR 伪中文乱码通常由模型"幻想"出随机 CJK 字符，
      每个字都不同，重复率极低。
    
    Returns:
        True 如果文本看起来是通顺的中文，False 如果是伪中文乱码
    """
    if not text:
        return False
    
    # 提取所有 CJK 字符
    cjk_chars = [ch for ch in text if '\u4e00' <= ch <= '\u9fff']
    if len(cjk_chars) < 4:
        return True  # CJK 太少，不判断
    
    unique_cjk = len(set(cjk_chars))
    total_cjk = len(cjk_chars)
    
    # 唯一字符率：真实文本通常 < 0.7，乱码通常 > 0.85
    uniqueness = unique_cjk / total_cjk if total_cjk > 0 else 1.0
    
    if uniqueness > 0.85 and total_cjk >= 8:
        return False  # 几乎所有字都不同 → 伪中文
    
    if uniqueness > 0.75 and total_cjk >= 15:
        return False  # 长文本且唯一率高 → 可疑
    
    # 额外检查：真实中文通常包含高频字
    high_freq_chars = set('的是不了一有人在上个大这来为和国我以要到他会就出对生能而子说时下过得自开部家机可方后成所分前然没法如经其现当于从者并部度实定物权加量都两体制当计还资应关因重些特线')
    high_freq_count = sum(1 for ch in cjk_chars if ch in high_freq_chars)
    high_freq_ratio = high_freq_count / total_cjk if total_cjk > 0 else 0
    
    # 真实文本高频字占比通常 > 15%
    if total_cjk >= 10 and high_freq_ratio < 0.08:
        return False
    
    return True


def _line_valid_char_ratio(line: str) -> float:
    """计算一行中有效字符的比例"""
    if not line:
        return 0.0
    valid = sum(1 for ch in line if (
        '\u4e00' <= ch <= '\u9fff' or
        '\u3400' <= ch <= '\u4dbf' or
        '\uf900' <= ch <= '\ufaff' or
        ch.isalnum() or
        ch in ' .,;:!?()（）[]【】""''、。，；：？！…—–-+=*/<>≤≥≠≈±×÷√{}%#@&^|~'
    ))
    return valid / len(line) if line else 0.0


def is_garbled_text(text: str, threshold: float = 0.35) -> bool:
    """
    判断文本是否为乱码（综合检测）。
    
    结合字符比例分析和中文连贯性检测。
    """
    if not text or len(text) < 5:
        return False
    
    # 使用 _line_valid_char_ratio 做快速检查
    ratio = _line_valid_char_ratio(text)
    if ratio < threshold:
        return True
    
    # CJK 字符多但不通顺 → 伪中文乱码
    cjk_count = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
    if cjk_count >= 10 and not is_coherent_chinese(text):
        return True
    
    # 拉丁字母占比过高且无常见单词模式 → 乱码
    latin = sum(1 for ch in text if ch.isascii() and ch.isalpha())
    total = len(text)
    if latin > total * 0.8 and len(text) > 20:
        common_patterns = re.findall(r'\b[a-z]{2,8}\b', text.lower())
        if len(common_patterns) < len(text) * 0.1:
            return True
    
    return False


# 判断是否为题号行
def is_question_line(line):
    """判断行是否为题号/标记行"""
    # 完全匹配题号模式
    if QUESTION_PREFIX_RE.match(line):
        return True
    # 匹配纯题号（如只有 "1" 或 "(1)"）
    pure_question_re = re.compile(
        r'^\s*(\d+|[a-zA-Z]|①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩)\s*$'
    )
    if pure_question_re.match(line):
        return True
    return False

# 判断是否为错题标记行（只有标记或主要是标记）
def is_wrong_marker_heavy(line):
    """判断行是否主要是错题标记"""
    stripped = line.strip()
    if not stripped:
        return False
    # 匹配仅包含标记的模式
    marker_only_re = re.compile(
        r'^(?:错|×|x|错误|答错|不对|不正确|✗|✘|❌|wrong|error|三角|红色|红叉|叉号|打叉|划掉|删除|取消)\s*$', 
        re.IGNORECASE
    )
    return bool(marker_only_re.match(stripped))

# ============================================================
# Pix2Text 单例管理
# ============================================================

_p2t_instance = None
_simple_lama_instance = None

def get_p2t():
    """
    获取 Pix2Text 单例实例。
    只初始化一次，避免重复加载模型导致的性能问题和识别不稳定。
    """
    global _p2t_instance
    if _p2t_instance is None:
        print("正在加载 Pix2Text 模型（仅首次）...")
        try:
            _p2t_instance = Pix2Text(
                analyzer_config=dict(
                    text_detector_config=dict(
                        box_thresh=0.3,
                        unclip_ratio=2.0
                    )
                )
            )
            print("Pix2Text 模型加载完成。")
        except Exception as e:
            print(f"Pix2Text 模型加载失败: {e}")
            raise
    return _p2t_instance


def _extract_text_from_result(result) -> str:
    """
    从 Pix2Text 返回结果中递归提取纯文本。
    统一处理 dict / list / str 等多种返回格式，避免编码损失。
    """
    if isinstance(result, str):
        return result
    elif isinstance(result, dict):
        # 优先取已知字段
        for key in ('text', 'content', 'line_texts', 'sentences'):
            if key in result:
                val = result[key]
                if isinstance(val, str):
                    return val
                elif isinstance(val, list):
                    return '\n'.join(_extract_text_from_result(v) for v in val)
        # 遍历所有值，收集文本
        parts = []
        for val in result.values():
            extracted = _extract_text_from_result(val)
            if extracted:
                parts.append(extracted)
        return '\n'.join(parts)
    elif isinstance(result, (list, tuple)):
        return '\n'.join(
            _extract_text_from_result(item) for item in result
        )
    elif result is not None:
        return str(result)
    return ''


def extract_text_from_image(image_path: Path) -> str:
    """
    从图片中提取文本，并清洗 OCR 乱码。
    
    Args:
        image_path: 图片文件路径
        
    Returns:
        清洗后的文本字符串
    """
    try:
        p2t = get_p2t()
        # resized_shape=1600：以更高分辨率处理，适合文字密集的试卷图片
        # 默认 768 对于 A4 试卷来说太低，容易漏识别小字和公式
        result = p2t.recognize(str(image_path), resized_shape=1600)
        raw_text = _extract_text_from_result(result)
        
        if not raw_text:
            print(f"  警告: 图片 {image_path.name} 未识别到文本")
            return ''
        
        # 清洗 OCR 噪声和乱码
        cleaned_text = clean_ocr_text(raw_text)
        
        # 检测是否为严重乱码（给出警告）
        if is_garbled_text(cleaned_text):
            print(f"  警告: 图片 {image_path.name} 识别结果疑似乱码，请检查图片质量")
        
        return cleaned_text
        
    except Exception as e:
        print(f"处理图片 {image_path} 时出错: {e}")
        return ""

def extract_question_number(line):
    """从行中提取题号"""
    match = QUESTION_PREFIX_RE.match(line)
    if match:
        # 查找第一个非空的捕获组（题号）
        for group in match.groups():
            if group:
                return group.strip()
    return None

def identify_wrong_answers(text):
    """
    识别错题：支持多种标记位置
    
    识别模式：
    1. 题号 + 错题标记（如：1. 错）
    2. 错题标记 + 题号 + 题干（如：错 1. xxx）
    3. 题号 + 包含标记的完整行（如：1. xxx 错）
    4. 标记在题目前的连续行
    """
    # 先清洗文本，去除 OCR 噪声
    text = clean_ocr_text(text)
    
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    wrong_questions = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 情况1: 题号后面跟着错题标记（标记在下一行）
        if is_question_line(line):
            question_num = extract_question_number(line)
            # 检查下一行是否是错题标记
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                if is_wrong_marker_heavy(next_line):
                    # 题号行 + 标记行
                    wrong_questions.append(f"{question_num} — {line}")
                    i += 2
                    continue
            # 检查当前行是否包含错题标记
            if WRONG_MARKER_RE.search(line):
                wrong_questions.append(f"{question_num} — {line}")
                i += 1
                continue
        
        # 情况2: 错题标记行，后面跟着题号
        if is_wrong_marker_heavy(line):
            # 检查是否紧跟着题号行
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                next_match = QUESTION_PREFIX_RE.match(next_line)
                if next_match:
                    question_num = extract_question_number(next_line)
                    wrong_questions.append(f"{question_num} — {next_line} (标记: {line})")
                    i += 2
                    continue
            # 单独一行错题标记，无法关联题号
            wrong_questions.append(f"未定位题号 — {line}")
            i += 1
            continue
        
        # 情况3: 当前行同时包含题号和错题标记
        if WRONG_MARKER_RE.search(line):
            question_num = extract_question_number(line)
            if question_num:
                wrong_questions.append(f"{question_num} — {line}")
            else:
                # 行中有标记但没有识别到题号
                wrong_questions.append(f"未定位题号 — {line}")
            i += 1
            continue
        
        i += 1
    
    # 去重（如果同一题被多次识别）
    seen = set()
    unique_wrong = []
    for q in wrong_questions:
        # 使用题号作为去重依据
        num_match = re.match(r'^(\d+|[a-zA-Z])', q)
        if num_match:
            num = num_match.group(1)
            if num not in seen:
                seen.add(num)
                unique_wrong.append(q)
        else:
            unique_wrong.append(q)
    
    return unique_wrong


def extract_all_questions(text):
    """提取所有题目块，按题号开始分组"""
    # 先清洗文本
    text = clean_ocr_text(text)
    
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    questions = []
    current_question = []

    for line in lines:
        if is_question_line(line):
            if current_question:
                questions.append(' '.join(current_question))
            current_question = [line]
        else:
            if current_question:
                current_question.append(line)

    if current_question:
        questions.append(' '.join(current_question))

    return questions


def create_word_document(items, output_path):
    """
    创建 Word 文档，设置中文字体避免打开时显示乱码。
    
    Args:
        items: 题目列表
        output_path: 输出文件路径
    """
    doc = Document()
    
    # ---- 设置默认字体为宋体，支持中文 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    # 设置中文字体（东亚字体）
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ---- 标题 ----
    heading = doc.add_heading('题目整理', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    if not items:
        p = doc.add_paragraph("未识别到题目。")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"共整理 {len(items)} 道题目", style='Normal')
        doc.add_paragraph('')  # 空行分隔
        
        for i, question in enumerate(items, 1):
            # 清理题目文本中可能残留的乱码
            clean_q = clean_ocr_text(question)
            
            heading = doc.add_heading(f'题目 {i}', level=1)
            for run in heading.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            para = doc.add_paragraph(clean_q)
            # 给每个题目段落也设置中文字体
            for run in para.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    try:
        doc.save(str(output_path))
        print(f"Word文档已保存到: {output_path}")
    except Exception as e:
        print(f"保存 Word 文档失败: {e}")
        raise

def main():
    parser = argparse.ArgumentParser(
        description='试卷截图处理工具：笔迹擦除 / OCR识别生成Word文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  # 只擦除笔迹，输出清洁图片
  python exam_scanner.py ./exam_screenshots --erase
  
  # OCR 识别所有题目，生成 Word 文档
  python exam_scanner.py ./exam_screenshots 错题集 --mode wrong
  python exam_scanner.py ./exam_screenshots 全部题目 --mode all
        '''
    )
    parser.add_argument('folder_path', help='包含试卷截图的文件夹路径')
    parser.add_argument('output_name', nargs='?', default=None,
                        help='生成的Word文档名称（不含扩展名），--erase 模式下不需要')
    parser.add_argument(
        '--erase', action='store_true',
        help='图片笔迹擦除模式：直接对图片做手写笔迹擦除，输出清洁图片（不生成Word文档）'
    )
    parser.add_argument(
        '--erase-backend',
        choices=['doc-image-tool', 'opencv', 'deeplabv3plus', 'dis', 'erasenet', 'lama', 'torchscript'],
        default='doc-image-tool',
        help=(
            '笔迹擦除后端：doc-image-tool=已接入的 Sauvola+K-Means 模型/工具管线；'
            'opencv=本地 mask+inpaint 规则管线；'
            'deeplabv3plus/dis/erasenet=第三方模型预测脚本或权重；'
            'lama=LaMa 生成式修复，默认尝试 IOPaint，也可配合 --model-command 和 {mask}；'
            'torchscript=直接加载 TorchScript/完整 PyTorch 模型'
        ),
    )
    parser.add_argument(
        '--model-repo',
        default=None,
        help='第三方模型仓库路径，如 HandWritingEraser-Pytorch / Handwriting-Removal-DIS / bdpan_erase_competition',
    )
    parser.add_argument(
        '--model-checkpoint',
        default=None,
        help='模型权重路径。TorchScript/完整模型可直接推理；普通 state_dict 建议配合 --model-command 使用',
    )
    parser.add_argument(
        '--model-command',
        default=None,
        help=(
            '第三方仓库预测命令模板，可用 {input} {staged_input} {mask} {output} '
            '{input_dir} {output_dir} {repo} {checkpoint} {device}。'
            '若输出灰度图则按 mask 白填；若输出彩色图则直接作为清洁图。'
            'lama 后端未设置该参数时会自动尝试 iopaint run。'
        ),
    )
    parser.add_argument(
        '--model-device',
        default='auto',
        help='模型推理设备：auto/cpu/cuda 等（默认: auto）',
    )
    parser.add_argument(
        '--mask-threshold',
        type=float,
        default=0.5,
        help='TorchScript mask 二值化阈值（默认: 0.5）',
    )
    parser.add_argument(
        '--mode', choices=['all', 'wrong'], default='all',
        help='整理模式：all=整理所有题目，wrong=仅整理错题（默认: all）'
    )

    args = parser.parse_args()

    folder_path = Path(args.folder_path)
    if not folder_path.exists() or not folder_path.is_dir():
        print(f"错误: 文件夹路径不存在: {folder_path}")
        sys.exit(1)

    # 支持的图片格式
    image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp'}

    # 收集所有图片文件
    image_files = sorted([
        f for f in folder_path.iterdir()
        if f.suffix.lower() in image_extensions
    ])

    if not image_files:
        print(f"错误: 文件夹 {folder_path} 中未找到支持的图片文件")
        print(f"支持的格式: {', '.join(image_extensions)}")
        sys.exit(1)

    # 笔迹擦除时跳过参考图「效果图」——只用于人工对照，不产生 *_cleaned 文件
    erase_targets = [f for f in image_files if not _is_reference_only_image(f)]

    # ============================================================
    # 笔迹擦除模式
    # ============================================================
    if args.erase:
        if not erase_targets:
            print(
                "错误: 文件夹中只有「效果图」参考图，没有可对齐 raw 。"
                "请放入手机拍摄的试卷照片（例如 raw.jpg）。"
            )
            sys.exit(1)

        output_dir = folder_path / "cleaned"
        output_dir.mkdir(exist_ok=True)
        model_repo = Path(args.model_repo).expanduser().resolve() if args.model_repo else None
        model_checkpoint = (
            Path(args.model_checkpoint).expanduser().resolve() if args.model_checkpoint else None
        )

        print(f"笔迹擦除模式")
        if len(erase_targets) < len(image_files):
            print(
                f"找到 {len(image_files)} 张图片，跳过参考图 "
                f"效果图.*（剩 {len(erase_targets)} 张待处理）"
            )
        else:
            print(f"找到 {len(erase_targets)} 张图片")
        print(f"输出目录: {output_dir}")
        print("-" * 50)

        success_count = 0
        for idx, file_path in enumerate(erase_targets, 1):
            output_path = output_dir / f"{file_path.stem}_cleaned{file_path.suffix}"
            print(f"[{idx}/{len(erase_targets)}] 擦除: {file_path.name}")

            try:
                remove_handwriting_from_image(
                    file_path,
                    output_path,
                    erase_backend=args.erase_backend,
                    model_repo=model_repo,
                    model_checkpoint=model_checkpoint,
                    model_command=args.model_command,
                    model_device=args.model_device,
                    mask_threshold=args.mask_threshold,
                )
                print(f"  → 已保存: {output_path.name}")
                success_count += 1
            except Exception as e:
                print(f"  ✗ 失败: {e}")

        print("-" * 50)
        print(f"完成: {success_count}/{len(erase_targets)} 张图片已处理")
        return

    # ============================================================
    # OCR 识别模式
    # ============================================================
    if not args.output_name:
        print("错误: OCR 模式下需要提供 output_name 参数")
        print("示例: python exam_scanner.py ./screenshots 我的错题集 --mode wrong")
        sys.exit(1)

    output_name = args.output_name
    if output_name.lower().endswith('.docx'):
        output_name = output_name[:-5]
    output_path = folder_path / f"{output_name}.docx"

    print(f"找到 {len(image_files)} 张图片")
    print(f"整理模式: {'仅错题' if args.mode == 'wrong' else '所有题目'}")
    print("-" * 50)

    all_items = []
    success_count = 0

    for idx, file_path in enumerate(image_files, 1):
        print(f"[{idx}/{len(image_files)}] 处理: {file_path.name}")
        text = extract_text_from_image(file_path)

        if not text:
            print(f"  跳过: 未识别到文本")
            continue

        if args.mode == 'all':
            items = extract_all_questions(text)
        else:
            items = identify_wrong_answers(text)

        if items:
            all_items.extend(items)
            print(f"  识别到 {len(items)} 道题目")
        else:
            print(f"  未识别到题目")

        success_count += 1

    print("-" * 50)
    print(f"处理完成: {success_count}/{len(image_files)} 张图片成功识别")
    print(f"共整理 {len(all_items)} 道题目")

    create_word_document(all_items, output_path)

if __name__ == "__main__":
    main()
